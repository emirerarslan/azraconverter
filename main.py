# -*- coding: utf-8 -*-

import sys
import os
import json
import hashlib
import csv
import re
import shutil
import subprocess
import tempfile
import threading
import traceback
import zipfile
from pathlib import Path
from urllib.error import URLError
from urllib.parse import urljoin, urlsplit
from urllib.request import Request, urlopen

# Dönüştürme kütüphaneleri açılışta değil, ilgili işlem seçildiğinde yüklenir.
# Böylece arayüz mümkün olan en kısa sürede görüntülenir.
OCR_AVAILABLE = None
_PDF_FONTS = None

from PySide6.QtCore import Qt, QObject, Signal, QThread, QTimer, QSettings, QUrl, QSize
from PySide6.QtGui import QFont, QIcon, QPixmap, QPainter, QPainterPath, QColor, QPen
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QGridLayout, QLabel, QPushButton, QFileDialog, QMessageBox, QFrame,
    QProgressBar, QSizePolicy, QSpacerItem, QDialog, QScrollArea,
    QTableWidget, QTableWidgetItem, QComboBox, QCheckBox, QButtonGroup, QLayout,
    QGraphicsDropShadowEffect
)
try:
    from PySide6.QtMultimedia import QAudioOutput, QMediaPlayer
    from PySide6.QtMultimediaWidgets import QVideoWidget
    QT_MULTIMEDIA_AVAILABLE = True
except (ImportError, ModuleNotFoundError):
    QAudioOutput = QMediaPlayer = QVideoWidget = None
    QT_MULTIMEDIA_AVAILABLE = False


APP_NAME = "ConverteR"
APP_VERSION = "1.1.21"
APP_ICON_FILE = "converter-new.ico"
UPDATE_CONFIG_FILE = "update_config.json"
DEFAULT_MANIFEST_URLS = [
    "https://github.com/emirerarslan/ConverteR/releases/latest/download/version.json",
    "https://api.github.com/repos/emirerarslan/ConverteR/releases/latest",
    "https://raw.githubusercontent.com/emirerarslan/ConverteR/main/updates/version.json",
]

PDF_EXTENSIONS = {".pdf"}
WORD_EXTENSIONS = {
    ".doc", ".docx", ".docm", ".dot", ".dotx", ".dotm",
    ".odt", ".rtf", ".txt",
}
SPREADSHEET_EXTENSIONS = {
    ".xls", ".xlsx", ".xlsm", ".xlsb", ".xlt", ".xltx", ".xltm",
    ".ods", ".csv", ".tsv",
}
SUPPORTED_EXTENSIONS = PDF_EXTENSIONS | WORD_EXTENSIONS | SPREADSHEET_EXTENSIONS


# Tema varlıkları tek noktada tanımlanır. İlk eşleşen dosya kullanıldığı için
# kaynak klasörü ve paketlenmiş uygulama aynı kurallarla çalışır.
THEME_ASSET_CANDIDATES = {
    "app_icon": (APP_ICON_FILE,),
    "azra_logo": ("azra-logo.png", "azra_gold_logo_real_transparent.png"),
    "rafine_logo": ("rafine-logo.jpg", "rafine-logo.png"),
    "emir_photo": ("emir-foto.png", "emir-foto.jpg", "emir-logo.jpg"),
    "emir_video": ("emir-video.mp4", "emir-video.mov", "emir-video.avi"),
    "turkish_flag": ("bayrak.jpeg", "setup_flag.bmp"),
    "emir_star": ("emir-yıldız.png", "emir-yildiz.png", "emir-yıldız.jpg"),
    "azra_converter_nav": ("Gold-donusturucu.png",),
    "azra_history_nav": ("silver-gecmis.png",),
    "azra_about_nav": ("bronz-hakkinda.png",),
    "rafine_converter_nav": ("dolar-donusturucu.png",),
    "rafine_history_nav": ("euro-gecmis.png",),
    "rafine_about_nav": ("frang-hakkinda.png",),
}

THEME_CONFIGS = {
    "azra": {
        "name": "Azra Mod",
        "logo": "azra_logo",
        "glow": QColor(235, 196, 111, 225),
        "halo": QColor(214, 177, 107, 50),
    },
    "rafine": {
        "name": "Rafine Mod",
        "logo": "rafine_logo",
        "glow": QColor(47, 196, 176, 225),
        "halo": QColor(47, 196, 176, 52),
    },
    "emir": {
        "name": "Emir Mod",
        "glow": QColor(210, 48, 62, 230),
        "halo": QColor(210, 48, 62, 58),
    },
}

THEME_STYLESHEETS = {
    "azra": """
        QMainWindow, QWidget#mainContent,
        QScrollArea#contentScroll,
        QScrollArea#contentScroll > QWidget > QWidget { background: #07110D; }
        QFrame#sidebar { background: #0C1712; border-right-color: #3F3521; }
        QLabel#pageTitle { color: #F7EBD0; }
        QLabel#eyebrow, QLabel#panelEyebrow, QLabel#sourceType { color: #D7B56D; }
        QPushButton#nav:checked { background: #1D2118; color: #E0BF74; border-left-color: #D7B56D; }
        QFrame#dropZone { background: #0D1612; border-color: #54472C; }
        QFrame#conversionPanel, QFrame#conversionCard { background: #101713; border-color: #3E382A; }
        QPushButton#selectButton, QPushButton#convertButton { background: #D7B56D; color: #10110D; }
        QPushButton#modeButton:checked { background: #D7B56D; color: #11120E; border-color: #E8CF96; }
    """,
    "rafine": """
        QMainWindow, QWidget#mainContent,
        QScrollArea#contentScroll,
        QScrollArea#contentScroll > QWidget > QWidget { background: #071214; }
        QFrame#sidebar { background: #0B1719; border-right-color: #1D4747; }
        QLabel#pageTitle { color: #EAF8F6; }
        QLabel#eyebrow, QLabel#panelEyebrow, QLabel#sourceType { color: #38C4B2; }
        QPushButton#nav:hover { background: #102526; border-color: #27625E; }
        QPushButton#nav:checked { background: #102827; color: #55D3C3; border-left-color: #38C4B2; }
        QFrame#dropZone { background: #0C181A; border-color: #2E5E5B; }
        QFrame#conversionPanel, QFrame#conversionCard { background: #0E1A1C; border-color: #28504E; }
        QComboBox#targetFormat { border-color: #2E7D75; }
        QPushButton#selectButton, QPushButton#convertButton { background: #38C4B2; color: #071311; }
        QProgressBar::chunk { background: #38C4B2; }
        QPushButton#modeToggle:hover, QPushButton#modeToggle:checked { background: #102827; color: #55D3C3; }
        QPushButton#modeButton:checked { background: #38C4B2; color: #071311; border-color: #6AE1D2; }
    """,
    "emir": """
        QMainWindow, QWidget#mainContent,
        QScrollArea#contentScroll,
        QScrollArea#contentScroll > QWidget > QWidget { background: #0D0C0E; }
        QFrame#sidebar { background: #121012; border-right-color: #4B2026; }
        QLabel#pageTitle { color: #FFF4EE; }
        QLabel#eyebrow, QLabel#panelEyebrow, QLabel#sourceType { color: #E3B768; }
        QPushButton#nav {
            background: #171215; color: #E8D7D4;
            border: 1px solid #A77B2E; border-radius: 10px;
        }
        QPushButton#nav:hover, QPushButton#nav:hover:checked {
            background: #35181F; color: #FFF6F1;
            border: 1px solid #FFD271;
        }
        QPushButton#nav:pressed { background: #4A1D27; border-color: #FFF0AF; }
        QPushButton#nav:checked { background: #171215; color: #FFF2EE; border: 1px solid #A77B2E; }
        QFrame#dropZone { background: #151114; border-color: #67313A; }
        QFrame#conversionPanel, QFrame#conversionCard { background: #171315; border-color: #4E2A31; }
        QComboBox#targetFormat { border-color: #7F3843; }
        QPushButton#selectButton, QPushButton#convertButton { background: #C92F40; color: #FFF8F4; }
        QPushButton#selectButton:hover, QPushButton#convertButton:hover { background: #DF4656; }
        QProgressBar::chunk { background: #C92F40; }
        QPushButton#modeToggle:hover, QPushButton#modeToggle:checked { background: #2A171B; color: #F0C5C9; }
        QPushButton#modeButton:checked { background: #C92F40; color: #FFF8F4; border-color: #ED6B77; }
    """,
}


def resource_path(name):
    base = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    return str(base / name)


def app_folder_path(name):
    """Paketlenmiş uygulamanın yanındaki, kullanıcı tarafından düzenlenebilir dosya."""
    base = Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else Path(__file__).resolve().parent
    return base / name


def theme_asset_path(asset_key):
    """Bir tema varlığını dış uygulama klasöründe veya paket içinde bulur."""
    for name in THEME_ASSET_CANDIDATES.get(asset_key, ()):
        external = app_folder_path(name)
        if external.exists():
            return str(external)
        bundled = Path(resource_path(name))
        if bundled.exists():
            return str(bundled)
    return ""


def version_key(version):
    parts = []
    for part in str(version).strip().lstrip("vV").split("."):
        digits = "".join(char for char in part if char.isdigit())
        parts.append(int(digits or 0))
    return tuple((parts + [0, 0, 0])[:3])


def load_update_config():
    config_path = app_folder_path(UPDATE_CONFIG_FILE)
    if not config_path.exists():
        return {}
    try:
        return json.loads(config_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


def update_manifest_urls(config):
    """Yapılandırılmış adresleri ve güvenli varsayılan yedekleri sırayla döndürür."""
    urls = []
    configured = config.get("manifest_urls", [])
    if isinstance(configured, str):
        configured = [configured]
    if isinstance(configured, list):
        urls.extend(clean_text(url) for url in configured)
    legacy_url = clean_text(config.get("manifest_url"))
    if legacy_url:
        urls.append(legacy_url)
    urls.extend(DEFAULT_MANIFEST_URLS)
    return list(dict.fromkeys(url for url in urls if url))


def normalise_update_manifest(payload, source_url):
    """Özel version.json ve GitHub Releases API yanıtlarını tek biçime getirir."""
    if clean_text(payload.get("tag_name")):
        version = clean_text(payload.get("tag_name")).lstrip("vV")
        assets = payload.get("assets") or []
        installer = next(
            (
                asset for asset in assets
                if clean_text(asset.get("name")).lower().endswith(".exe")
                and "setup" in clean_text(asset.get("name")).lower()
            ),
            None,
        )
        update_package = next(
            (
                asset for asset in assets
                if clean_text(asset.get("name")).lower().endswith(".zip")
                and "update" in clean_text(asset.get("name")).lower()
            ),
            None,
        )
        download_url = clean_text((installer or {}).get("browser_download_url"))
        digest = clean_text((installer or {}).get("digest"))
        sha256 = digest.split(":", 1)[1] if digest.lower().startswith("sha256:") else ""
        package_url = clean_text((update_package or {}).get("browser_download_url"))
        package_digest = clean_text((update_package or {}).get("digest"))
        package_sha256 = (
            package_digest.split(":", 1)[1]
            if package_digest.lower().startswith("sha256:") else ""
        )
        return {
            "version": version,
            "download_url": download_url,
            "sha256": sha256,
            "package_url": package_url,
            "package_sha256": package_sha256,
            "notes": clean_text(payload.get("body"))[:300],
        }

    download_url = clean_text(payload.get("download_url"))
    if download_url:
        download_url = urljoin(source_url, download_url)
    package_url = clean_text(payload.get("package_url"))
    if package_url:
        package_url = urljoin(source_url, package_url)
    return {
        "version": clean_text(payload.get("version")),
        "download_url": download_url,
        "sha256": clean_text(payload.get("sha256")),
        "package_url": package_url,
        "package_sha256": clean_text(payload.get("package_sha256")),
        "notes": clean_text(payload.get("notes")),
    }


def extract_update_package(package_path, destination):
    """Doğrulanmış ZIP'i yol geçişi ve sembolik bağlantı saldırılarına karşı açar."""
    package_path = Path(package_path).resolve()
    destination = Path(destination).resolve()
    destination.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(package_path) as archive:
        for member in archive.infolist():
            member_path = Path(member.filename.replace("\\", "/"))
            if member_path.is_absolute() or ".." in member_path.parts:
                raise RuntimeError("Güncelleme paketi güvenli olmayan bir yol içeriyor.")
            # Unix ZIP'lerinde sembolik bağlantı türü üst 16 bitte tutulur.
            if (member.external_attr >> 16) & 0o170000 == 0o120000:
                raise RuntimeError("Güncelleme paketi sembolik bağlantı içeriyor.")
            target = (destination / member_path).resolve()
            if destination != target and destination not in target.parents:
                raise RuntimeError("Güncelleme paketi hedef klasör dışına çıkıyor.")
        archive.extractall(destination)

    executable = next(destination.rglob(f"{APP_NAME}.exe"), None)
    if executable is None:
        raise RuntimeError("Güncelleme paketinde uygulama dosyası bulunamadı.")
    return executable.parent




def register_turkish_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    candidates = [
        (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
        (r"C:\Windows\Fonts\segoeui.ttf", r"C:\Windows\Fonts\segoeuib.ttf"),
        (r"C:\Windows\Fonts\calibri.ttf", r"C:\Windows\Fonts\calibrib.ttf"),
    ]

    for regular, bold in candidates:
        if os.path.exists(regular):
            try:
                pdfmetrics.registerFont(TTFont("AzraFont", regular))
                if os.path.exists(bold):
                    pdfmetrics.registerFont(TTFont("AzraFontBold", bold))
                else:
                    pdfmetrics.registerFont(TTFont("AzraFontBold", regular))
                return "AzraFont", "AzraFontBold"
            except Exception:
                pass

    return "Helvetica", "Helvetica-Bold"


def get_pdf_fonts():
    global _PDF_FONTS
    if _PDF_FONTS is None:
        _PDF_FONTS = register_turkish_fonts()
    return _PDF_FONTS


def get_ocr_dependencies():
    global OCR_AVAILABLE

    if OCR_AVAILABLE is False:
        return None

    try:
        import fitz
        import pytesseract
        from PIL import Image
        OCR_AVAILABLE = True
        return fitz, pytesseract, Image
    except ImportError:
        OCR_AVAILABLE = False
        return None


def clean_text(value):
    if value is None:
        return ""
    return str(value).replace("\x00", "").strip()


def ocr_text_to_columns(line):
    """Düz PDF metnindeki belirgin boşluk ve ayraçları Excel sütunlarına böler."""
    text = clean_text(line)
    if not text:
        return []
    if "\t" in text:
        parts = re.split(r"\t+", text)
    elif "|" in text:
        parts = text.split("|")
    else:
        parts = re.split(r"\s{2,}", text)
    values = [clean_text(part) for part in parts if clean_text(part)]
    return values or [text]


def unique_output(path):
    path = Path(path)
    if not path.exists():
        return path

    i = 1
    while True:
        candidate = path.with_name(f"{path.stem}_{i}{path.suffix}")
        if not candidate.exists():
            return candidate
        i += 1


def _ps_quote(value):
    """Bir yolu PowerShell tek tırnaklı sabitinde güvenli hale getirir."""
    return str(Path(value).resolve()).replace("'", "''")


def update_result_path():
    """Güncelleme yardımcısının sonraki açılışa bıraktığı sonucu döndürür."""
    local_data = Path(os.environ.get("LOCALAPPDATA", Path.home() / "AppData" / "Local"))
    return local_data / "Azra Converter" / "update-result.json"


def write_update_result(status, version, message=""):
    """Yönetici olarak çalışan yardımcı ile arayüz arasında kalıcı durum kaydı."""
    path = update_result_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps({"status": status, "version": version, "message": message}, ensure_ascii=False),
        encoding="utf-8",
    )


def consume_update_result():
    path = update_result_path()
    try:
        result = json.loads(path.read_text(encoding="utf-8-sig"))
        path.unlink(missing_ok=True)
        return result if isinstance(result, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def _run_powershell_automation(script, failure_message):
    """Office COM otomasyonunu görünür pencere açmadan çalıştırır."""
    startupinfo = subprocess.STARTUPINFO()
    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    startupinfo.wShowWindow = subprocess.SW_HIDE
    result = subprocess.run(
        [
            "powershell", "-NoProfile", "-NonInteractive",
            "-ExecutionPolicy", "Bypass", "-Command", script,
        ],
        capture_output=True,
        text=True,
        startupinfo=startupinfo,
        check=False,
    )
    if result.returncode != 0:
        detail = clean_text(result.stderr) or clean_text(result.stdout)
        raise RuntimeError(
            failure_message + (f"\n\nAyrıntı: {detail}" if detail else "")
        )


def find_libreoffice():
    candidates = [
        os.environ.get("LIBREOFFICE_PATH", ""),
        shutil.which("soffice") or "",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    return next((item for item in candidates if item and Path(item).exists()), None)


def ensure_spreadsheet_dependency_compatibility():
    """OpenPyXL'i yeni NumPy sürümleriyle güvenli biçimde yükler.

    Bazı NumPy dağıtımları ``numpy.short`` takma adını kaldırabiliyor. Eski
    OpenPyXL sürümleri bu adı modül yüklenirken kullandığından XLSX içeren tüm
    dönüşümler başlamadan kesiliyordu. NumPy isteğe bağlıdır; kurulu değilse
    OpenPyXL kendi standart sayı türleriyle çalışmaya devam eder.
    """
    try:
        import numpy
    except ImportError:
        return

    aliases = {
        "short": "int16",
        "ushort": "uint16",
        "intc": "int32",
        "uintc": "uint32",
        "longlong": "int64",
        "ulonglong": "uint64",
    }
    for legacy_name, replacement_name in aliases.items():
        if not hasattr(numpy, legacy_name) and hasattr(numpy, replacement_name):
            setattr(numpy, legacy_name, getattr(numpy, replacement_name))


def libreoffice_convert(src, out, target_format):
    """LibreOffice ile güvenli bir geçici klasörde biçim dönüştürür."""
    soffice = find_libreoffice()
    if not soffice:
        raise RuntimeError("LibreOffice kurulu değil.")

    filters = {
        "pdf": "pdf",
        "docx": "docx:Office Open XML Text",
        "xlsx": "xlsx:Calc MS Excel 2007 XML",
    }
    if target_format not in filters:
        raise ValueError(f"LibreOffice hedef biçimi desteklenmiyor: {target_format}")

    src = Path(src).resolve()
    out = Path(out).resolve()
    with tempfile.TemporaryDirectory(prefix="AzraConverter-") as temp_dir:
        result = subprocess.run(
            [
                soffice, "--headless", "--nologo", "--nodefault", "--nolockcheck",
                "--convert-to", filters[target_format], "--outdir", temp_dir, str(src),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        candidates = list(Path(temp_dir).glob(f"*.{target_format}"))
        if result.returncode != 0 or not candidates:
            detail = clean_text(result.stderr) or clean_text(result.stdout)
            raise RuntimeError(
                "LibreOffice dönüşümü tamamlayamadı."
                + (f"\n\nAyrıntı: {detail}" if detail else "")
            )
        out.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(candidates[0], out)
    return out


def word_to_docx_with_microsoft_word(src, out):
    script = (
        "$ErrorActionPreference='Stop'; $word=$null; $document=$null; "
        "$word=New-Object -ComObject Word.Application; $word.Visible=$false; "
        "$word.DisplayAlerts=0; try { "
        f"$document=$word.Documents.Open('{_ps_quote(src)}',$false,$true); "
        f"$document.SaveAs2('{_ps_quote(out)}',16); "
        "} finally { if ($document) {$document.Close($false)}; "
        "if ($word) {$word.Quit()} }"
    )
    _run_powershell_automation(script, "Microsoft Word belgeyi DOCX biçimine çeviremedi.")
    if not Path(out).exists():
        raise RuntimeError("Microsoft Word DOCX çıktısını oluşturamadı.")


def excel_to_xlsx_with_microsoft_excel(src, out):
    script = (
        "$ErrorActionPreference='Stop'; $excel=$null; $book=$null; "
        "$excel=New-Object -ComObject Excel.Application; $excel.Visible=$false; "
        "$excel.DisplayAlerts=$false; try { "
        f"$book=$excel.Workbooks.Open('{_ps_quote(src)}',0,$true); "
        f"$book.SaveAs('{_ps_quote(out)}',51); "
        "} finally { if ($book) {$book.Close($false)}; "
        "if ($excel) {$excel.Quit()} }"
    )
    _run_powershell_automation(script, "Microsoft Excel dosyayı XLSX biçimine çeviremedi.")
    if not Path(out).exists():
        raise RuntimeError("Microsoft Excel XLSX çıktısını oluşturamadı.")


def _normalise_word_document(src, temp_dir):
    src = Path(src)
    if src.suffix.lower() == ".docx":
        return src
    target = Path(temp_dir) / f"{src.stem}.docx"
    try:
        word_to_docx_with_microsoft_word(src, target)
    except Exception as word_error:
        try:
            libreoffice_convert(src, target, "docx")
        except Exception as libreoffice_error:
            raise RuntimeError(
                f"{src.suffix.upper()} belgesi açılamadı. Microsoft Word veya "
                f"LibreOffice gereklidir.\n\nWord: {word_error}\nLibreOffice: {libreoffice_error}"
            )
    return target


def _normalise_spreadsheet(src, temp_dir):
    """Yaygın tablo biçimlerini kayıpsız işlem için XLSX'e normalleştirir."""
    ensure_spreadsheet_dependency_compatibility()
    from openpyxl import Workbook

    src = Path(src)
    if src.suffix.lower() in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return src

    target = Path(temp_dir) / f"{src.stem}.xlsx"
    if src.suffix.lower() in {".csv", ".tsv"}:
        raw = src.read_bytes()
        decoded = None
        for encoding in ("utf-8-sig", "utf-8", "cp1254", "latin-1"):
            try:
                decoded = raw.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if decoded is None:
            raise RuntimeError("CSV/TSV dosyasının karakter kodlaması okunamadı.")

        delimiter = "\t"
        if src.suffix.lower() == ".csv":
            try:
                delimiter = csv.Sniffer().sniff(
                    decoded[:8192], delimiters=",;\t|"
                ).delimiter
            except csv.Error:
                delimiter = ";" if decoded.count(";") > decoded.count(",") else ","
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Veriler"
        for row in csv.reader(decoded.splitlines(), delimiter=delimiter):
            sheet.append(row)
        workbook.save(target)
        return target

    try:
        excel_to_xlsx_with_microsoft_excel(src, target)
    except Exception as excel_error:
        try:
            libreoffice_convert(src, target, "xlsx")
        except Exception as libreoffice_error:
            raise RuntimeError(
                f"{src.suffix.upper()} çalışma kitabı açılamadı. Microsoft Excel "
                f"veya LibreOffice gereklidir.\n\nExcel: {excel_error}\n"
                f"LibreOffice: {libreoffice_error}"
            )
    return target



def find_tesseract():
    candidates = [
        os.environ.get("TESSERACT_CMD", ""),
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]

    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate

    try:
        import shutil
        return shutil.which("tesseract")
    except Exception:
        return None


def smart_ocr_page(page):
    """Taranmış PDF sayfasını OCR ile okuyup sütunları koordinatlardan çıkarır."""
    dependencies = get_ocr_dependencies()
    if not dependencies:
        raise RuntimeError(
            "OCR paketleri kurulu değil. "
            "python -m pip install pytesseract pillow pymupdf"
        )

    fitz, pytesseract, Image = dependencies

    tess = find_tesseract()

    if not tess:
        raise RuntimeError(
            r"Tesseract bulunamadı: C:\Program Files\Tesseract-OCR\tesseract.exe"
        )

    pytesseract.pytesseract.tesseract_cmd = tess

    pix = page.get_pixmap(
        matrix=fitz.Matrix(2.5, 2.5),
        alpha=False
    )

    image = Image.frombytes(
        "RGB",
        [pix.width, pix.height],
        pix.samples
    )

    try:
        data = pytesseract.image_to_data(
            image,
            lang="tur+eng",
            config="--psm 6",
            output_type=pytesseract.Output.DICT,
        )
    except Exception:
        data = pytesseract.image_to_data(
            image,
            lang="eng",
            config="--psm 6",
            output_type=pytesseract.Output.DICT,
        )

    words = []

    for i, raw in enumerate(data.get("text", [])):
        text = str(raw).strip()

        try:
            confidence = float(data["conf"][i])
        except Exception:
            confidence = -1

        if not text or confidence < 15:
            continue

        left = int(data["left"][i])
        top = int(data["top"][i])
        width = int(data["width"][i])
        height = int(data["height"][i])

        words.append({
            "text": text,
            "left": left,
            "right": left + width,
            "top": top,
            "height": height,
            "center_y": top + height / 2,
        })

    # Kelimeleri yatay satırlara ayır.
    lines = []

    for word in sorted(
        words,
        key=lambda x: (x["top"], x["left"])
    ):
        target = None

        for line in lines:
            tolerance = max(
                8,
                int(line["avg_height"] * 0.65)
            )

            if abs(
                word["center_y"] - line["center_y"]
            ) <= tolerance:
                target = line
                break

        if target is None:
            lines.append({
                "center_y": word["center_y"],
                "avg_height": max(word["height"], 1),
                "words": [word],
            })
        else:
            target["words"].append(word)

            target["center_y"] = sum(
                x["center_y"]
                for x in target["words"]
            ) / len(target["words"])

            target["avg_height"] = sum(
                x["height"]
                for x in target["words"]
            ) / len(target["words"])

    rows = []

    for line in sorted(
        lines,
        key=lambda x: x["center_y"]
    ):
        row_words = sorted(
            line["words"],
            key=lambda x: x["left"]
        )

        cells = []

        for word in row_words:
            if not cells:
                cells.append({
                    "text": word["text"],
                    "right": word["right"],
                })
                continue

            gap = word["left"] - cells[-1]["right"]

            # Büyük yatay boşluk = yeni sütun.
            threshold = max(
                28,
                int(line["avg_height"] * 1.6)
            )

            if gap >= threshold:
                cells.append({
                    "text": word["text"],
                    "right": word["right"],
                })
            else:
                cells[-1]["text"] += " " + word["text"]
                cells[-1]["right"] = word["right"]

        if cells:
            rows.append([
                clean_text(cell["text"])
                for cell in cells
            ])

    return rows


def pdf_to_excel(src, progress, separate_pages=True):
    """PDF içeriğini Excel'e aktarır.

    ``separate_pages`` kapalı olduğunda PDF sayfalarının satırları tek bir
    çalışma sayfasında sırayla birleştirilir.
    """
    import pdfplumber
    src = Path(src)
    out = unique_output(
        src.with_name(src.stem + "_Excel.xlsx")
    )

    ensure_spreadsheet_dependency_compatibility()
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    wb = Workbook()
    first_sheet = True
    combined_sheet = wb.active if not separate_pages else None
    if combined_sheet is not None:
        combined_sheet.title = "Veriler"
    combined_row_index = 1

    with pdfplumber.open(str(src)) as pdf:
        total = max(len(pdf.pages), 1)

        for page_no, page in enumerate(
            pdf.pages,
            start=1
        ):
            if separate_pages:
                ws = wb.active if first_sheet else wb.create_sheet()
                ws.title = f"Sayfa {page_no}"
                first_sheet = False
                row_index = 1
            else:
                ws = combined_sheet
                row_index = combined_row_index

            tables = page.extract_tables()

            if tables:
                for table in tables:
                    for row in table:
                        values = [
                            clean_text(x)
                            for x in (row or [])
                        ]

                        if values:
                            for col_index, value in enumerate(
                                values,
                                start=1
                            ):
                                ws.cell(
                                    row=row_index,
                                    column=col_index,
                                    value=value
                                )

                            row_index += 1

                    row_index += 1

            else:
                text = page.extract_text() or ""

                if text.strip():
                    rows = [
                        ocr_text_to_columns(line)
                        for line in text.splitlines()
                        if clean_text(line)
                    ]
                else:
                    # Taranmış PDF -> koordinat tabanlı OCR.
                    dependencies = get_ocr_dependencies()
                    if not dependencies:
                        raise RuntimeError("OCR paketleri kurulu değil.")
                    fitz, _, _ = dependencies
                    with fitz.open(str(src)) as ocr_doc:
                        rows = smart_ocr_page(
                            ocr_doc[page_no - 1]
                        )

                for values in rows:
                    for col_index, value in enumerate(
                        values,
                        start=1
                    ):
                        ws.cell(
                            row=row_index,
                            column=col_index,
                            value=clean_text(value)
                        )
                    row_index += 1

            if ws.max_row:
                for cell in ws[1]:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(
                        vertical="top"
                    )

            for col in ws.columns:
                letter = col[0].column_letter

                max_len = max(
                    (
                        len(clean_text(cell.value))
                        for cell in col
                    ),
                    default=0
                )

                ws.column_dimensions[
                    letter
                ].width = min(
                    max(max_len + 2, 10),
                    45
                )

            if not separate_pages:
                combined_row_index = row_index

            progress.emit(
                int(page_no / total * 100)
            )

    wb.save(str(out))
    return out


def pdf_page_count(src):
    """PDF'in sayfa sayısını hızlıca ve metin/OCR işlemi yapmadan döndürür."""
    from pypdf import PdfReader

    reader = PdfReader(str(src))
    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                raise RuntimeError("Şifreli PDF'in sayfaları okunamadı.")
        except Exception as exc:
            raise RuntimeError("Şifreli PDF'in sayfaları okunamadı.") from exc
    return len(reader.pages)


def remove_pdf_pages(src, pages_to_remove, progress):
    """Seçilen 1 tabanlı sayfa numaralarını çıkarıp yeni bir PDF oluşturur."""
    from pypdf import PdfReader, PdfWriter

    src = Path(src)
    out = unique_output(src.with_name(src.stem + "_Sayfalari_Silinmis.pdf"))
    reader = PdfReader(str(src))
    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                raise RuntimeError("Şifreli PDF'in sayfaları silinemiyor.")
        except Exception as exc:
            raise RuntimeError("Şifreli PDF'in sayfaları silinemiyor.") from exc

    total = len(reader.pages)
    remove_set = {int(page) for page in pages_to_remove}
    if not remove_set:
        raise RuntimeError("Silmek için en az bir sayfa seçin.")
    if any(page < 1 or page > total for page in remove_set):
        raise RuntimeError("Geçersiz PDF sayfa numarası seçildi.")
    if len(remove_set) >= total:
        raise RuntimeError("PDF'in tüm sayfaları silinemez. En az bir sayfa kalmalı.")

    writer = PdfWriter()
    for index, page in enumerate(reader.pages, start=1):
        if index not in remove_set:
            writer.add_page(page)
        progress.emit(int(index / total * 95))

    metadata = reader.metadata
    if metadata:
        writer.add_metadata({
            str(key): str(value)
            for key, value in metadata.items()
            if value is not None
        })
    with open(out, "wb") as output_file:
        writer.write(output_file)
    progress.emit(100)
    return out



def pdf_to_word(src, progress):
    import pdfplumber
    from docx import Document
    from docx.shared import Pt
    src = Path(src)
    out = unique_output(
        src.with_name(src.stem + "_Word.docx")
    )

    # Word'ün PDF yeniden akış motoru metin kutularını, görselleri ve sayfa
    # düzenini yerleşik çıkarıcıdan daha başarılı korur. Kullanılabiliyorsa önce
    # onu dener, başarısız olursa aşağıdaki OCR/tablo yoluna geçeriz.
    try:
        progress.emit(5)
        word_to_docx_with_microsoft_word(src, out)
        progress.emit(100)
        return out
    except Exception:
        if out.exists():
            out.unlink()

    doc = Document()

    with pdfplumber.open(str(src)) as pdf:
        total = max(len(pdf.pages), 1)

        for page_no, page in enumerate(
            pdf.pages,
            start=1
        ):
            if page_no > 1:
                doc.add_page_break()

            tables = page.extract_tables()

            if tables:
                for table_data in tables:
                    rows = [
                        [
                            clean_text(x)
                            for x in (row or [])
                        ]
                        for row in table_data
                    ]

                    max_cols = max(
                        (
                            len(row)
                            for row in rows
                        ),
                        default=0
                    )

                    if rows and max_cols:
                        table = doc.add_table(
                            rows=len(rows),
                            cols=max_cols
                        )

                        table.style = "Table Grid"

                        for r, row in enumerate(rows):
                            for c in range(max_cols):
                                table.cell(
                                    r,
                                    c
                                ).text = (
                                    row[c]
                                    if c < len(row)
                                    else ""
                                )

                        doc.add_paragraph()

            else:
                text = page.extract_text() or ""

                if text.strip():
                    lines = [
                        clean_text(line)
                        for line in text.splitlines()
                        if clean_text(line)
                    ]
                else:
                    dependencies = get_ocr_dependencies()
                    if not dependencies:
                        raise RuntimeError("OCR paketleri kurulu değil.")
                    fitz, _, _ = dependencies
                    with fitz.open(str(src)) as ocr_doc:
                        rows = smart_ocr_page(
                            ocr_doc[page_no - 1]
                        )

                    lines = [
                        "    ".join(row)
                        for row in rows
                        if row
                    ]

                for line in lines:
                    paragraph = doc.add_paragraph(line)
                    paragraph.paragraph_format.space_after = Pt(4)

            progress.emit(
                int(page_no / total * 100)
            )

    doc.save(str(out))
    return out



def make_pdf_styles():
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.styles import ParagraphStyle
    pdf_font, pdf_font_bold = get_pdf_fonts()
    return {
        "title": ParagraphStyle(
            "AzraTitle",
            fontName=pdf_font_bold,
            fontSize=15,
            leading=19,
            alignment=TA_LEFT,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "AzraBody",
            fontName=pdf_font,
            fontSize=9,
            leading=12,
            alignment=TA_LEFT,
        ),
        "small": ParagraphStyle(
            "AzraSmall",
            fontName=pdf_font,
            fontSize=7.5,
            leading=9,
        ),
    }


def excel_to_pdf_with_microsoft_excel(src, out):
    """Excel'in baskı alanı, grafik ve sayfa ayarlarını koruyarak PDF üretir."""
    script = (
        "$ErrorActionPreference='Stop'; $excel=$null; $book=$null; "
        "$excel=New-Object -ComObject Excel.Application; $excel.Visible=$false; "
        "$excel.DisplayAlerts=$false; try { "
        f"$book=$excel.Workbooks.Open('{_ps_quote(src)}',0,$true); "
        f"$book.ExportAsFixedFormat(0,'{_ps_quote(out)}'); "
        "} finally { if ($book) {$book.Close($false)}; "
        "if ($excel) {$excel.Quit()} }"
    )
    _run_powershell_automation(script, "Microsoft Excel PDF çıktısını oluşturamadı.")
    if not Path(out).exists():
        raise RuntimeError("Microsoft Excel PDF çıktısını oluşturamadı.")


def excel_to_pdf(src, progress):
    ensure_spreadsheet_dependency_compatibility()
    from openpyxl import load_workbook
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, PageBreak

    src = Path(src)
    out = unique_output(src.with_name(src.stem + "_PDF.pdf"))

    progress.emit(5)
    try:
        excel_to_pdf_with_microsoft_excel(src, out)
        progress.emit(100)
        return out
    except Exception:
        if out.exists():
            out.unlink()

    try:
        libreoffice_convert(src, out, "pdf")
        progress.emit(100)
        return out
    except Exception:
        if out.exists():
            out.unlink()

    if src.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        raise RuntimeError(
            f"{src.suffix.upper()} dosyasını PDF'e dönüştürmek için Microsoft "
            "Excel veya LibreOffice kurulu olmalıdır."
        )

    wb = load_workbook(str(src), data_only=True)
    sheet_names = wb.sheetnames
    styles = make_pdf_styles()
    pdf_font, pdf_font_bold = get_pdf_fonts()

    doc = SimpleDocTemplate(
        str(out),
        pagesize=landscape(A4),
        rightMargin=10 * mm,
        leftMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )

    story = []
    total = max(len(sheet_names), 1)

    for index, sheet_name in enumerate(sheet_names, start=1):
        ws = wb[sheet_name]

        story.append(Paragraph(clean_text(sheet_name), styles["title"]))

        data = []
        for row in ws.iter_rows(values_only=True):
            values = [clean_text(v) for v in row]
            while values and values[-1] == "":
                values.pop()
            if values:
                data.append(values)

        if data:
            max_cols = max(len(row) for row in data)
            data = [row + [""] * (max_cols - len(row)) for row in data]

            usable_width = landscape(A4)[0] - 20 * mm
            col_width = usable_width / max_cols

            table = Table(
                [[Paragraph(v.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["small"])
                  for v in row] for row in data],
                colWidths=[col_width] * max_cols,
                repeatRows=1,
            )

            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D6B16B")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTNAME", (0, 0), (-1, 0), pdf_font_bold),
                ("FONTNAME", (0, 1), (-1, -1), pdf_font),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))

            story.append(table)
        else:
            story.append(Paragraph("Bu sayfada veri bulunamadı.", styles["body"]))

        if index < total:
            story.append(PageBreak())

        progress.emit(int(index / total * 100))

    doc.build(story)
    return out


def word_to_pdf_with_microsoft_word(src, out):
    """Word'ün kendi PDF dışa aktarmasıyla düzeni yüksek sadakatle korur."""
    script = (
        "$ErrorActionPreference='Stop'; $word=$null; $document=$null; "
        "$word=New-Object -ComObject Word.Application; $word.Visible=$false; "
        "$word.DisplayAlerts=0; "
        "try { "
        f"$document=$word.Documents.Open('{_ps_quote(src)}',$false,$true); "
        f"$document.ExportAsFixedFormat('{_ps_quote(out)}',17); "
        "} finally { "
        "if ($document) {$document.Close($false)}; if ($word) {$word.Quit()} "
        "}"
    )
    _run_powershell_automation(script, "Microsoft Word PDF çıktısını oluşturamadı.")
    if not Path(out).exists():
        raise RuntimeError("Microsoft Word PDF çıktısını oluşturamadı.")


def word_to_pdf(src, progress):
    from docx import Document
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    src = Path(src)
    out = unique_output(src.with_name(src.stem + "_PDF.pdf"))

    # Önce Word, sonra LibreOffice: bu iki yol resim, üstbilgi, dipnot,
    # grafik ve sayfa sonlarını yerleşik metin tabanlı yoldan daha iyi korur.
    progress.emit(5)
    try:
        word_to_pdf_with_microsoft_word(src, out)
        progress.emit(100)
        return out
    except Exception:
        if out.exists():
            out.unlink()

    try:
        libreoffice_convert(src, out, "pdf")
        progress.emit(100)
        return out
    except Exception:
        if out.exists():
            out.unlink()

    if src.suffix.lower() != ".docx":
        raise RuntimeError(
            f"{src.suffix.upper()} belgesini PDF'e dönüştürmek için Microsoft "
            "Word veya LibreOffice kurulu olmalıdır."
        )

    document = Document(str(src))
    styles = make_pdf_styles()
    pdf_font, pdf_font_bold = get_pdf_fonts()

    doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )

    story = []
    blocks = max(len(document.paragraphs) + len(document.tables), 1)
    done = 0

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if text:
            style = styles["body"]
            if paragraph.style and paragraph.style.name:
                name = paragraph.style.name.lower()
                if "title" in name or "heading" in name:
                    style = styles["title"]

            safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, style))
            story.append(Spacer(1, 3 * mm))

        done += 1
        progress.emit(int(done / blocks * 100))

    for table_data in document.tables:
        rows = []

        for row in table_data.rows:
            values = []
            for cell in row.cells:
                values.append(clean_text(cell.text))
            rows.append(values)

        if rows:
            max_cols = max(len(r) for r in rows)
            rows = [r + [""] * (max_cols - len(r)) for r in rows]

            usable_width = A4[0] - 36 * mm
            col_width = usable_width / max_cols

            table = Table(
                [[Paragraph(v.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["small"])
                  for v in row] for row in rows],
                colWidths=[col_width] * max_cols,
                repeatRows=1,
            )

            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D6B16B")),
                ("FONTNAME", (0, 0), (-1, 0), pdf_font_bold),
                ("FONTNAME", (0, 1), (-1, -1), pdf_font),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))

            story.append(Spacer(1, 4 * mm))
            story.append(table)
            story.append(Spacer(1, 4 * mm))

        done += 1
        progress.emit(int(done / blocks * 100))

    if not story:
        story.append(Paragraph("Belgede dönüştürülebilecek içerik bulunamadı.", styles["body"]))

    doc.build(story)
    return out


def word_to_excel(src, progress):
    """Word metnini ve her tabloyu düzenlenebilir Excel sayfalarına aktarır."""
    from docx import Document
    ensure_spreadsheet_dependency_compatibility()
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    src = Path(src)
    out = unique_output(src.with_name(src.stem + "_Excel.xlsx"))
    with tempfile.TemporaryDirectory(prefix="AzraConverter-") as temp_dir:
        normalised = _normalise_word_document(src, temp_dir)
        document = Document(str(normalised))
        workbook = Workbook()
        text_sheet = workbook.active
        text_sheet.title = "Belge Metni"
        text_sheet.append(["Sıra", "Tür", "İçerik"])

        paragraphs = [p for p in document.paragraphs if clean_text(p.text)]
        total = max(len(paragraphs) + len(document.tables), 1)
        done = 0
        for number, paragraph in enumerate(paragraphs, start=1):
            style_name = clean_text(getattr(paragraph.style, "name", ""))
            kind = "Başlık" if any(
                marker in style_name.lower() for marker in ("title", "heading", "başlık")
            ) else "Metin"
            text_sheet.append([number, kind, clean_text(paragraph.text)])
            done += 1
            progress.emit(int(done / total * 90))

        for table_number, source_table in enumerate(document.tables, start=1):
            sheet = workbook.create_sheet(f"Tablo {table_number}")
            for row in source_table.rows:
                sheet.append([clean_text(cell.text) for cell in row.cells])
            done += 1
            progress.emit(int(done / total * 90))

        header_fill = PatternFill("solid", fgColor="D6B16B")
        for sheet in workbook.worksheets:
            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions
            for cell in sheet[1]:
                cell.font = Font(bold=True)
                cell.fill = header_fill
                cell.alignment = Alignment(vertical="top", wrap_text=True)
            for column in sheet.columns:
                letter = column[0].column_letter
                max_length = max((len(clean_text(cell.value)) for cell in column), default=0)
                sheet.column_dimensions[letter].width = min(max(max_length + 2, 10), 60)
                for cell in column:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)

        workbook.save(out)
    progress.emit(100)
    return out


def excel_to_word(src, progress):
    """Çalışma kitabındaki tüm sayfaları düzenlenebilir Word tablolarına aktarır."""
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches, Pt
    ensure_spreadsheet_dependency_compatibility()
    from openpyxl import load_workbook

    src = Path(src)
    out = unique_output(src.with_name(src.stem + "_Word.docx"))
    with tempfile.TemporaryDirectory(prefix="AzraConverter-") as temp_dir:
        normalised = _normalise_spreadsheet(src, temp_dir)
        values_book = load_workbook(str(normalised), data_only=True, read_only=True)
        formulas_book = load_workbook(str(normalised), data_only=False, read_only=True)
        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(9)

        total = max(len(values_book.sheetnames), 1)
        needs_landscape = any(values_book[name].max_column > 6 for name in values_book.sheetnames)
        if needs_landscape:
            section = document.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            section.left_margin = section.right_margin = Inches(0.45)

        for index, sheet_name in enumerate(values_book.sheetnames, start=1):
            value_sheet = values_book[sheet_name]
            formula_sheet = formulas_book[sheet_name]
            if index > 1:
                document.add_page_break()
            document.add_heading(clean_text(sheet_name), level=1)

            rows = []
            for value_row, formula_row in zip(
                value_sheet.iter_rows(values_only=True),
                formula_sheet.iter_rows(values_only=True),
            ):
                values = []
                for cached, formula in zip(value_row, formula_row):
                    value = cached if cached is not None else formula
                    values.append(clean_text(value))
                while values and values[-1] == "":
                    values.pop()
                if values:
                    rows.append(values)

            if rows:
                max_columns = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=max_columns)
                table.style = "Table Grid"
                for row_index, values in enumerate(rows):
                    for column_index in range(max_columns):
                        table.cell(row_index, column_index).text = (
                            values[column_index] if column_index < len(values) else ""
                        )
                for cell in table.rows[0].cells:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
            else:
                document.add_paragraph("Bu sayfada veri bulunamadı.")

            progress.emit(int(index / total * 95))

        document.save(out)
        values_book.close()
        formulas_book.close()
    progress.emit(100)
    return out


class UpdateWorker(QObject):
    finished = Signal(object)
    error = Signal(str)

    def run(self):
        config = load_update_config()
        manifest = None
        errors = []
        for manifest_url in update_manifest_urls(config):
            try:
                request = Request(
                    manifest_url,
                    headers={
                        "User-Agent": f"AzraConverter-Updater/{APP_VERSION}",
                        "Accept": "application/vnd.github+json, application/json",
                        "Cache-Control": "no-cache",
                    },
                )
                with urlopen(request, timeout=12) as response:
                    payload = json.loads(response.read().decode("utf-8-sig"))
                candidate = normalise_update_manifest(payload, manifest_url)
                if candidate["version"]:
                    manifest = candidate
                    break
                errors.append(f"{urlsplit(manifest_url).netloc}: sürüm bilgisi yok")
            except (OSError, URLError, json.JSONDecodeError, UnicodeDecodeError) as exc:
                errors.append(f"{urlsplit(manifest_url).netloc}: {exc}")

        if manifest is None:
            detail = "; ".join(errors[:3])
            self.error.emit(
                "Güncelleme servislerine ulaşılamadı. İnternet bağlantısını kontrol "
                f"edip tekrar deneyin.{f' Ayrıntı: {detail}' if detail else ''}"
            )
            return

        latest = manifest["version"]
        if not latest:
            self.error.emit("Güncelleme bilgisinde sürüm numarası bulunamadı.")
            return

        self.finished.emit({
            "version": latest,
            "is_new": version_key(latest) > version_key(APP_VERSION),
            "download_url": manifest["download_url"],
            "sha256": manifest["sha256"],
            "package_url": manifest.get("package_url", ""),
            "package_sha256": manifest.get("package_sha256", ""),
            "notes": manifest["notes"],
        })


class UpdateDownloadWorker(QObject):
    progress = Signal(int)
    finished = Signal(str)
    error = Signal(str)

    def __init__(self, url, version, expected_sha256=""):
        super().__init__()
        self.url = url
        self.version = version
        self.expected_sha256 = expected_sha256.lower()

    def run(self):
        try:
            update_folder = Path(tempfile.gettempdir()) / "AzraConverterUpdates"
            update_folder.mkdir(parents=True, exist_ok=True)
            name = Path(urlsplit(self.url).path).name or f"AzraConverter-{self.version}-Setup.exe"
            target = update_folder / name
            partial = target.with_suffix(target.suffix + ".download")

            request = Request(
                self.url,
                headers={"User-Agent": "AzraConverter-Updater"},
            )
            with urlopen(request, timeout=30) as response, open(partial, "wb") as output:
                total = int(response.headers.get("Content-Length", 0))
                received = 0
                while True:
                    chunk = response.read(1024 * 256)
                    if not chunk:
                        break
                    output.write(chunk)
                    received += len(chunk)
                    if total:
                        self.progress.emit(min(100, int(received / total * 100)))

            if not partial.exists() or partial.stat().st_size == 0:
                raise RuntimeError("Güncelleme paketi indirilemedi.")
            if total and received != total:
                raise RuntimeError("Güncelleme paketi eksik indirildi. Lütfen tekrar deneyin.")

            actual_sha256 = hashlib.sha256(partial.read_bytes()).hexdigest().lower()
            if self.expected_sha256 and actual_sha256 != self.expected_sha256:
                raise RuntimeError("Güncelleme paketi doğrulanamadı. Lütfen tekrar deneyin.")

            os.replace(partial, target)
            self.progress.emit(100)
            self.finished.emit(str(target))
        except Exception as exc:
            self.error.emit(str(exc) or "Güncelleme paketi indirilemedi.")


class ConversionCancelled(Exception):
    """Kullanıcının güvenli bir kontrol noktasında dönüşümü iptal etmesi."""


class CancellableProgress:
    def __init__(self, worker):
        self.worker = worker

    def emit(self, value):
        self.worker.check_cancelled()
        self.worker.progress.emit(value)


class ConverterWorker(QObject):
    finished = Signal(str)
    error = Signal(str)
    progress = Signal(int)
    cancelled = Signal()

    def __init__(
        self, mode, source, pdf_excel_separate_pages=True, pages_to_remove=None
    ):
        super().__init__()
        self.mode = mode
        self.source = source
        self.pdf_excel_separate_pages = pdf_excel_separate_pages
        self.pages_to_remove = pages_to_remove or []
        self._cancel_event = threading.Event()

    def cancel(self):
        self._cancel_event.set()

    def check_cancelled(self):
        if self._cancel_event.is_set():
            raise ConversionCancelled()

    def run(self):
        try:
            self.check_cancelled()
            progress = CancellableProgress(self)
            if self.mode == "pdf_excel":
                result = pdf_to_excel(
                    self.source, progress,
                    separate_pages=self.pdf_excel_separate_pages,
                )
            elif self.mode == "pdf_delete_pages":
                result = remove_pdf_pages(
                    self.source, self.pages_to_remove, progress
                )
            elif self.mode == "pdf_word":
                result = pdf_to_word(self.source, progress)
            elif self.mode == "excel_pdf":
                result = excel_to_pdf(self.source, progress)
            elif self.mode == "word_pdf":
                result = word_to_pdf(self.source, progress)
            elif self.mode == "word_excel":
                result = word_to_excel(self.source, progress)
            elif self.mode == "excel_word":
                result = excel_to_word(self.source, progress)
            else:
                raise ValueError("Geçersiz dönüşüm seçildi.")

            self.check_cancelled()
            self.finished.emit(str(result))
        except ConversionCancelled:
            self.cancelled.emit()
        except Exception as exc:
            # Kullanıcıya hata izinin tamamı yerine doğrudan çözüm odaklı mesajı göster.
            self.error.emit(str(exc) or traceback.format_exc())



class DropZone(QFrame):
    fileDropped = Signal(str)
    clicked = Signal()

    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        self.setCursor(Qt.PointingHandCursor)
        self.setObjectName("dropZone")
        self._glow_phase = 0.0
        self._glow_color = QColor(235, 196, 111, 225)
        self._halo_color = QColor(214, 177, 107, 50)
        self._border_timer = QTimer(self)
        self._border_timer.timeout.connect(self._advance_border_glow)
        self._border_timer.start(42)

        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignCenter)
        layout.setSpacing(8)

        self.icon = QLabel("DOSYA")
        self.icon.setAlignment(Qt.AlignCenter)
        self.icon.setObjectName("dropIcon")

        self.title = QLabel("DOSYAYI BURAYA BIRAK")
        self.title.setAlignment(Qt.AlignCenter)
        self.title.setObjectName("dropTitle")
        self.title.setWordWrap(True)

        self.sub = QLabel("PDF | DOC/DOCX | XLS/XLSX | ODT/ODS | CSV/RTF")
        self.sub.setAlignment(Qt.AlignCenter)
        self.sub.setObjectName("dropSub")
        self.sub.setWordWrap(True)

        layout.addWidget(self.icon)
        layout.addWidget(self.title)
        layout.addWidget(self.sub)

    def _advance_border_glow(self):
        """Altın şerit ışığını çerçevenin etrafında kesintisiz ilerletir."""
        if not self._should_show_border():
            return
        self._glow_phase = (self._glow_phase + 1.35) % 72
        self.update()

    def _should_show_border(self):
        return bool(
            self.property("selected")
            or self.property("dragging")
            or self.underMouse()
        )

    def set_glow_colors(self, glow, halo):
        self._glow_color = QColor(glow)
        self._halo_color = QColor(halo)
        self.update()

    def paintEvent(self, event):
        super().paintEvent(event)
        if not self._should_show_border():
            return

        # Stil sayfasının ince çerçevesinin üzerine iki katmanlı, hareketli bir
        # LED şeridi çizilir. Kesiklerin kayması, ışığın çerçeveyi dolaştığı
        # izlenimini verir; parlak katman ise altın yıldız parıltısı oluşturur.
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        rect = self.rect().adjusted(3, 3, -3, -3)

        halo = QPen(self._halo_color, 5)
        halo.setStyle(Qt.DashLine)
        halo.setDashPattern([1.2, 11.0])
        halo.setDashOffset(-self._glow_phase)
        halo.setCapStyle(Qt.RoundCap)
        painter.setPen(halo)
        painter.drawRoundedRect(rect, 14, 14)

        lights = QPen(self._glow_color, 1.7)
        lights.setStyle(Qt.DashLine)
        lights.setDashPattern([1.0, 12.5])
        lights.setDashOffset(-self._glow_phase)
        lights.setCapStyle(Qt.RoundCap)
        painter.setPen(lights)
        painter.drawRoundedRect(rect, 14, 14)
        painter.end()

    def enterEvent(self, event):
        self.update()
        super().enterEvent(event)

    def leaveEvent(self, event):
        self.update()
        super().leaveEvent(event)

    def show_file(self, path):
        path = Path(path)
        try:
            size = path.stat().st_size
            size_text = (
                f"{size / (1024 * 1024):.1f} MB"
                if size >= 1024 * 1024 else f"{max(1, size // 1024)} KB"
            )
        except OSError:
            size_text = ""
        self.setProperty("selected", True)
        self.style().unpolish(self)
        self.style().polish(self)
        self.icon.setText(path.suffix.lstrip(".").upper() or "DOSYA")
        self.title.setText(path.name)
        self.sub.setText(
            f"{size_text}  |  Değiştirmek için tıklayın veya yeni dosya bırakın"
            if size_text else "Değiştirmek için tıklayın veya yeni dosya bırakın"
        )

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.clicked.emit()
        super().mousePressEvent(event)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
            self.setProperty("dragging", True)
            self.style().unpolish(self)
            self.style().polish(self)
            self.update()

    def dragLeaveEvent(self, event):
        self.setProperty("dragging", False)
        self.style().unpolish(self)
        self.style().polish(self)
        self.update()
        super().dragLeaveEvent(event)

    def dropEvent(self, event):
        self.setProperty("dragging", False)
        self.style().unpolish(self)
        self.style().polish(self)
        self.update()
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            if path:
                self.fileDropped.emit(path)
        event.acceptProposedAction()


class ConversionProgressDialog(QDialog):
    cancelRequested = Signal()

    def __init__(self, file_name, parent=None):
        super().__init__(parent)
        self._running = True
        self._cancel_requested = False
        self.setWindowTitle("Dönüştürülüyor")
        self.setModal(True)
        self.setWindowModality(Qt.ApplicationModal)
        self.setFixedSize(470, 245)
        self.setStyleSheet("""
            QDialog { background: #0F0F10; border: 1px solid #4A3C25; }
            QLabel#progressTitle {
                color: #D6B16B; font-size: 21px; font-weight: 800;
            }
            QLabel#progressFile {
                color: #F2E8D3; font-size: 14px; font-weight: 650;
            }
            QLabel#progressHint { color: #9A9286; font-size: 11px; }
            QProgressBar {
                background: #1A1917; color: #E7C87F;
                border: 1px solid #3B3429; border-radius: 7px;
                min-height: 24px; text-align: center; font-weight: 700;
            }
            QProgressBar::chunk { background: #D6B16B; border-radius: 6px; }
            QPushButton {
                background: #1B1916; color: #D6B16B;
                border: 1px solid #6B5633; border-radius: 8px;
                min-height: 36px; font-weight: 800;
            }
            QPushButton:hover { background: #D6B16B; color: #11100E; }
            QPushButton:disabled { color: #75684F; border-color: #39342C; }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(28, 24, 28, 22)
        layout.setSpacing(13)

        title = QLabel("DÖNÜŞTÜRÜLÜYOR")
        title.setObjectName("progressTitle")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        self.file_label = QLabel(file_name)
        self.file_label.setObjectName("progressFile")
        self.file_label.setAlignment(Qt.AlignCenter)
        self.file_label.setWordWrap(True)
        layout.addWidget(self.file_label)

        self.progress = QProgressBar()
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        layout.addWidget(self.progress)

        self.hint = QLabel("Lütfen işlem tamamlanana kadar bekleyin.")
        self.hint.setObjectName("progressHint")
        self.hint.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.hint)

        self.cancel_button = QPushButton("İPTAL ET")
        self.cancel_button.clicked.connect(self.request_cancel)
        layout.addWidget(self.cancel_button)

    def set_progress(self, value):
        self.progress.setValue(value)

    def request_cancel(self):
        if not self._running or self._cancel_requested:
            return
        self._cancel_requested = True
        self.cancel_button.setEnabled(False)
        self.cancel_button.setText("İPTAL EDİLİYOR...")
        self.hint.setText("İşlem güvenli biçimde durduruluyor, lütfen bekleyin.")
        self.cancelRequested.emit()

    def finish(self):
        self._running = False
        self.close()

    def closeEvent(self, event):
        if self._running:
            self.request_cancel()
            event.ignore()
            return
        super().closeEvent(event)


class PdfPageDeletionDialog(QDialog):
    """PDF sayfalarını silme seçimini anlaşılır bir ızgarada sunar."""

    def __init__(self, page_count, parent=None):
        super().__init__(parent)
        self._page_count = page_count
        self._page_checks = []
        self.setWindowTitle("PDF Sayfalarını Sil")
        self.setModal(True)
        self.resize(520, 510)
        self.setStyleSheet("""
            QDialog { background: #0F0F10; border: 1px solid #4A3C25; }
            QLabel#deleteTitle { color: #D6B16B; font-size: 20px; font-weight: 800; }
            QLabel#deleteHint, QLabel#deleteSummary { color: #BBB3A7; font-size: 12px; }
            QScrollArea { background: #151413; border: 1px solid #3B3429; border-radius: 9px; }
            QCheckBox {
                color: #E6DED0; background: #1B1916; border: 1px solid #3D362B;
                border-radius: 7px; padding: 8px; font-weight: 700;
            }
            QCheckBox:hover { border-color: #B99552; background: #242016; }
            QCheckBox:checked { color: #FFCE88; background: #352217; border-color: #D17B3E; }
            QPushButton {
                background: #1B1916; color: #D6B16B; border: 1px solid #6B5633;
                border-radius: 8px; min-height: 34px; padding: 0 13px; font-weight: 800;
            }
            QPushButton:hover { background: #D6B16B; color: #11100E; }
            QPushButton#deleteConfirm { background: #C95B3C; color: #FFF8F2; border-color: #E38B64; }
            QPushButton#deleteConfirm:hover { background: #E16A47; }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 22, 24, 20)
        layout.setSpacing(11)

        title = QLabel("SİLİNECEK SAYFALARI SEÇİN")
        title.setObjectName("deleteTitle")
        layout.addWidget(title)

        hint = QLabel(
            "İşaretlenen sayfalar silinir. İşaretlenmeyen sayfalar PDF'te kalır."
        )
        hint.setObjectName("deleteHint")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        actions = QHBoxLayout()
        select_all = QPushButton("TÜMÜNÜ SEÇ")
        select_all.clicked.connect(lambda: self._set_all_checked(True))
        clear_all = QPushButton("SEÇİMİ TEMİZLE")
        clear_all.clicked.connect(lambda: self._set_all_checked(False))
        actions.addWidget(select_all)
        actions.addWidget(clear_all)
        actions.addStretch(1)
        layout.addLayout(actions)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        page_list = QWidget()
        grid = QGridLayout(page_list)
        grid.setContentsMargins(12, 12, 12, 12)
        grid.setSpacing(8)
        for page_number in range(1, page_count + 1):
            check = QCheckBox(f"Sayfa {page_number}")
            check.stateChanged.connect(self._update_summary)
            self._page_checks.append(check)
            grid.addWidget(check, (page_number - 1) // 3, (page_number - 1) % 3)
        scroll.setWidget(page_list)
        layout.addWidget(scroll, 1)

        self.summary = QLabel()
        self.summary.setObjectName("deleteSummary")
        layout.addWidget(self.summary)
        self._update_summary()

        buttons = QHBoxLayout()
        buttons.addStretch(1)
        cancel = QPushButton("İPTAL")
        cancel.clicked.connect(self.reject)
        confirm = QPushButton("SEÇİLEN SAYFALARI SİL")
        confirm.setObjectName("deleteConfirm")
        confirm.clicked.connect(self._confirm)
        buttons.addWidget(cancel)
        buttons.addWidget(confirm)
        layout.addLayout(buttons)

    def selected_pages(self):
        return [
            index for index, check in enumerate(self._page_checks, start=1)
            if check.isChecked()
        ]

    def _set_all_checked(self, checked):
        for check in self._page_checks:
            check.setChecked(checked)
        self._update_summary()

    def _update_summary(self, *_args):
        deleted = len(self.selected_pages())
        remaining = self._page_count - deleted
        self.summary.setText(
            f"Silinecek: {deleted} sayfa  •  Kalacak: {remaining} sayfa"
        )

    def _confirm(self):
        deleted = len(self.selected_pages())
        if not deleted:
            QMessageBox.information(self, "Sayfa seçin", "Silmek için en az bir sayfa işaretleyin.")
            return
        if deleted == self._page_count:
            QMessageBox.warning(self, "En az bir sayfa kalsın", "PDF'in tüm sayfaları silinemez.")
            return
        self.accept()


class NavButton(QPushButton):
    """Temaya göre özgün görsel veya sade Emir gezinti düğmesi."""

    IMAGE_ASSETS = {
        "azra": ("azra_converter_nav", "azra_history_nav", "azra_about_nav"),
        "rafine": ("rafine_converter_nav", "rafine_history_nav", "rafine_about_nav"),
    }
    ACCENT_COLORS = {
        "azra": QColor("#F2CA62"),
        "rafine": QColor("#D9C2E8"),
    }

    def __init__(self, text, nav_index):
        super().__init__(text)
        self._label = text.strip()
        self._nav_index = nav_index
        self._theme_key = "emir"
        self._nav_pixmap = QPixmap()
        self._scaled_nav_pixmap = QPixmap()
        self._scaled_nav_cache_key = None
        self._hover_glow = None
        self.setCheckable(True)
        self.setMinimumHeight(46)
        self.setAccessibleName(self._label)

    def set_theme(self, theme_key):
        self._theme_key = theme_key
        asset_keys = self.IMAGE_ASSETS.get(theme_key)
        asset_path = theme_asset_path(asset_keys[self._nav_index]) if asset_keys else ""
        self._nav_pixmap = QPixmap(asset_path) if asset_path else QPixmap()
        self._scaled_nav_pixmap = QPixmap()
        self._scaled_nav_cache_key = None
        uses_image = not self._nav_pixmap.isNull()
        self.setText("" if uses_image else f"  {self._label}")
        self.setIcon(QIcon(theme_asset_path("emir_star")) if theme_key == "emir" else QIcon())
        self.setIconSize(QSize(18, 18))
        self.setMinimumHeight(82 if uses_image else 46)
        self.setProperty("imageNav", uses_image)
        if self._hover_glow is not None:
            self._hover_glow.setEnabled(theme_key == "emir" and self.underMouse())
        self.style().unpolish(self)
        self.style().polish(self)
        self.update()

    def paintEvent(self, event):
        super().paintEvent(event)
        if self._nav_pixmap.isNull():
            if self._theme_key == "emir":
                painter = QPainter(self)
                painter.setRenderHint(QPainter.Antialiasing)
                painter.setPen(QPen(QColor("#5A421B"), 1))
                painter.drawRoundedRect(self.rect().adjusted(4, 4, -5, -5), 7, 7)
                painter.setPen(QPen(QColor(244, 199, 91, 105), 0.7))
                painter.drawRoundedRect(self.rect().adjusted(6, 6, -7, -7), 5, 5)
                painter.end()
            return

        # Butonlar Windows ekran ölçeklemesinde (ör. %125 / %150) bulanık
        # görünmesin diye görseli mantıksal değil, fiziksel piksel boyutunda
        # üretiriz. Sonuç tekrar kullanılacağından her çizimde yeniden
        # örnekleme yapılmaz.
        ratio = self.devicePixelRatioF()
        cache_key = (self.width(), self.height(), round(ratio * 100))
        if cache_key != self._scaled_nav_cache_key:
            pixel_size = QSize(
                max(1, round(self.width() * ratio)),
                max(1, round(self.height() * ratio)),
            )
            self._scaled_nav_pixmap = self._nav_pixmap.scaled(
                pixel_size, Qt.KeepAspectRatio, Qt.SmoothTransformation
            )
            self._scaled_nav_pixmap.setDevicePixelRatio(ratio)
            self._scaled_nav_cache_key = cache_key

        painter = QPainter(self)
        painter.setRenderHint(QPainter.SmoothPixmapTransform)
        pixmap = self._scaled_nav_pixmap
        pixmap_width = round(pixmap.width() / ratio)
        pixmap_height = round(pixmap.height() / ratio)
        x = (self.width() - pixmap_width) // 2
        y = (self.height() - pixmap_height) // 2
        painter.drawPixmap(x, y, pixmap)
        # Görsel sekmelerde sabit bir dış çerçeve kullanılmaz. Vurgu yalnızca
        # imleç gerçekten butonun üzerindeyken görünür.
        if self.underMouse():
            accent = self.ACCENT_COLORS.get(self._theme_key, QColor("#D6B16B"))
            accent.setAlpha(185)
            pen = QPen(accent, 1.25)
            painter.setPen(pen)
            painter.drawRoundedRect(self.rect().adjusted(2, 2, -3, -3), 10, 10)
        painter.end()

    def enterEvent(self, event):
        if self._theme_key == "emir":
            if self._hover_glow is None:
                self._hover_glow = QGraphicsDropShadowEffect(self)
                self._hover_glow.setBlurRadius(28)
                self._hover_glow.setOffset(0, 0)
                self._hover_glow.setColor(QColor(239, 72, 91, 225))
                self.setGraphicsEffect(self._hover_glow)
            self._hover_glow.setEnabled(True)
        self.update()
        super().enterEvent(event)

    def leaveEvent(self, event):
        if self._theme_key == "emir":
            if self._hover_glow is not None:
                self._hover_glow.setEnabled(False)
        self.update()
        super().leaveEvent(event)


class ModeButton(QPushButton):
    def __init__(self, text, mode_key):
        super().__init__(text)
        self.mode_key = mode_key
        self.setObjectName("modeButton")
        self.setCheckable(True)
        self.setMinimumHeight(31)


class ConversionCard(QFrame):
    def __init__(self, title, subtitle, parent=None):
        super().__init__(parent)
        self.setObjectName("conversionCard")
        self.setCursor(Qt.PointingHandCursor)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 16, 18, 16)
        layout.setSpacing(5)

        self.icon = QLabel(">>")
        self.icon.setObjectName("cardIcon")
        self.icon.setAlignment(Qt.AlignLeft)

        self.title = QLabel(title)
        self.title.setObjectName("cardTitle")

        self.subtitle = QLabel(subtitle)
        self.subtitle.setObjectName("cardSubtitle")
        self.subtitle.setWordWrap(True)

        layout.addWidget(self.icon)
        layout.addWidget(self.title)
        layout.addWidget(self.subtitle)

        self.button = QPushButton("BAŞLAT")
        self.button.setObjectName("cardButton")
        self.button.setMinimumHeight(34)
        layout.addWidget(self.button)

    def set_available(self, available):
        """Kartı, seçilen kaynak dosya türü için kullanılabilir yapar."""
        self.setEnabled(available)
        self.setCursor(
            Qt.PointingHandCursor if available else Qt.ArrowCursor
        )


class FirstRunThemeDialog(QDialog):
    """Yeni kurulumda bir kez gösterilen, karanlık tema seçim ekranı."""

    OPTIONS = (
        ("azra", "AZRA MOD", "Altın  •  Koyu  •  Premium", "#E8C56F", "#FFF0B5", "#241E10"),
        ("rafine", "RAFİNE MOD", "Finans  •  Modern  •  Profesyonel", "#35D5C0", "#8DFFF0", "#0C2927"),
        ("emir", "EMİR MOD", "Güçlü  •  Koyu  •  Türk kırmızısı", "#E43C50", "#FF8794", "#321017"),
    )

    def __init__(self, parent=None):
        super().__init__(parent)
        self.selected_theme = None
        self.setWindowTitle(f"{APP_NAME} — Tema Seçimi")
        self.setModal(True)
        self.setMinimumSize(650, 500)
        self.resize(760, 570)

        icon_path = theme_asset_path("app_icon")
        if icon_path:
            self.setWindowIcon(QIcon(icon_path))

        self.setStyleSheet("""
            QDialog {
                background: #07080A;
                color: #F8F4EA;
                font-family: "Palatino Linotype", Palatino, serif;
            }
            QLabel#onboardingBrand {
                color: #C9A95D;
                font-size: 11px;
                font-weight: 700;
                letter-spacing: 3px;
            }
            QLabel#onboardingTitle {
                color: #FFFDF7;
                font-size: 27px;
                font-weight: 700;
            }
            QLabel#onboardingText {
                color: #AAA7A1;
                font-family: "Segoe UI";
                font-size: 12px;
            }
            QLabel#onboardingNote {
                background: #0E1014;
                color: #BEBBB4;
                border: 1px solid #282B31;
                border-radius: 8px;
                padding: 10px 14px;
                font-family: "Segoe UI";
                font-size: 11px;
            }
            QFrame#onboardingLine {
                background: #B8974E;
                border: none;
            }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(52, 38, 52, 34)
        layout.setSpacing(13)

        brand = QLabel("CONVERTER")
        brand.setObjectName("onboardingBrand")
        brand.setAlignment(Qt.AlignCenter)
        layout.addWidget(brand)

        title = QLabel("Hangi temayı seçmek istiyorsunuz?")
        title.setObjectName("onboardingTitle")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        text = QLabel("Çalışma alanınız için başlangıç görünümünü belirleyin.")
        text.setObjectName("onboardingText")
        text.setAlignment(Qt.AlignCenter)
        layout.addWidget(text)

        line = QFrame()
        line.setObjectName("onboardingLine")
        line.setFixedHeight(2)
        layout.addWidget(line)
        layout.addSpacing(5)

        for key, name, description, accent, bright, hover in self.OPTIONS:
            button = QPushButton(f"{name}     —     {description}")
            button.setCursor(Qt.PointingHandCursor)
            button.setMinimumHeight(70)
            button.setStyleSheet(f"""
                QPushButton {{
                    background: #101216;
                    color: #F8F6F0;
                    border: 1px solid {accent};
                    border-left: 7px solid {accent};
                    border-radius: 10px;
                    padding: 12px 24px;
                    text-align: left;
                    font-family: "Palatino Linotype", Palatino, serif;
                    font-size: 14px;
                    font-weight: 700;
                }}
                QPushButton:hover {{
                    background: {hover};
                    color: #FFFFFF;
                    border: 2px solid {bright};
                    border-left: 8px solid {bright};
                }}
                QPushButton:pressed {{
                    background: #050607;
                }}
            """)
            glow = QGraphicsDropShadowEffect(button)
            glow.setBlurRadius(23)
            glow.setOffset(0, 0)
            glow.setColor(QColor(accent))
            button.setGraphicsEffect(glow)
            button.clicked.connect(lambda _checked=False, mode=key: self._choose(mode))
            layout.addWidget(button)

        layout.addStretch(1)
        note = QLabel("NOT  —  Seçtiğiniz tema daha sonra sol menüdeki Modlar bölümünden değiştirilebilir.")
        note.setObjectName("onboardingNote")
        note.setAlignment(Qt.AlignCenter)
        note.setWordWrap(True)
        layout.addWidget(note)

    def _choose(self, mode_key):
        self.selected_theme = mode_key
        self.accept()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.source_file = None
        self.thread = None
        self.worker = None
        self.progress_dialog = None
        self._active_mode = None
        self.current_page = "converter"
        self.active_theme = None

        self.setWindowTitle(APP_NAME)
        # İçerik kaydırılabildiğinden pencere daha küçük boyutlarda da
        # kullanılabilir. Son kullanılan boyut ve konum sonraki açılışta korunur.
        self.setMinimumSize(640, 480)
        self._window_settings = QSettings(APP_NAME, APP_NAME)
        saved_geometry = self._window_settings.value("window_geometry")
        if saved_geometry:
            self.restoreGeometry(saved_geometry)
        else:
            self.resize(1200, 780)

        icon_path = theme_asset_path("app_icon")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        self._base_stylesheet = """
            QMainWindow {
                background: #0A0A0B;
            }
            QScrollArea#contentScroll,
            QScrollArea#contentScroll > QWidget > QWidget {
                background: #0A0A0B;
                border: none;
            }
            QScrollArea#sidebarScroll {
                background: #101011;
                border: none;
            }
            QWidget {
                color: #F3F1EC;
                font-family: "Segoe UI";
            }
            QFrame#sidebar {
                background: #101011;
                border-right: 1px solid #252321;
            }
            QLabel#logoText {
                color: #D6B16B;
                font-size: 17px;
                font-weight: 600;
            }
            QLabel#version {
                color: #66615A;
                font-size: 11px;
            }
            QLabel#pageTitle {
                color: #F7F4ED;
                font-size: 29px;
                font-weight: 750;
            }
            QLabel#pageSubtitle {
                color: #85817A;
                font-size: 13px;
            }
            QLabel#eyebrow {
                color: #D6B16B;
                font-size: 11px;
                font-weight: 700;
                letter-spacing: 1px;
            }
            QPushButton#nav {
                background: transparent;
                color: #8D8982;
                border: none;
                border-radius: 9px;
                text-align: left;
                padding: 0 14px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton#nav:hover {
                background: #24211B;
                border: 1px solid #4A3C25;
                color: #E9E5DC;
            }
            QPushButton#nav:pressed {
                background: #2A241A;
                color: #E2C27C;
            }
            QPushButton#nav:checked {
                background: #211E18;
                color: #D6B16B;
                border-left: 2px solid #D6B16B;
            }
            QPushButton#nav[imageNav="true"],
            QPushButton#nav[imageNav="true"]:hover,
            QPushButton#nav[imageNav="true"]:pressed,
            QPushButton#nav[imageNav="true"]:checked {
                background: transparent;
                border: none;
                padding: 0;
            }
            QFrame#modePanel {
                background: #151413;
                border: 1px solid #302C27;
                border-radius: 10px;
            }
            QPushButton#modeToggle {
                background: transparent;
                color: #C6BDAF;
                border: none;
                border-radius: 7px;
                min-height: 32px;
                text-align: left;
                padding: 0 9px;
                font-size: 11px;
                font-weight: 800;
                letter-spacing: 1px;
            }
            QPushButton#modeToggle:hover,
            QPushButton#modeToggle:checked {
                background: #211F1B;
                color: #E4C57E;
            }
            QPushButton#modeButton {
                background: #1A1917;
                color: #A7A198;
                border: 1px solid #34312C;
                border-radius: 7px;
                text-align: left;
                padding: 0 10px;
                font-size: 11px;
                font-weight: 700;
            }
            QPushButton#modeButton:hover {
                color: #F0EAE0;
                border-color: #6B5633;
            }
            QFrame#dropZone {
                background: #111112;
                border: none;
                border-radius: 16px;
                min-height: 210px;
            }
            QFrame#dropZone:hover,
            QFrame#dropZone[dragging="true"] {
                background: #15130F;
                border: 1px dashed #D6B16B;
            }
            QFrame#dropZone[selected="true"] {
                background: #17140F;
                border: 2px solid #D6B16B;
            }
            QLabel#dropIcon {
                color: #D6B16B;
                font-size: 35px;
                font-weight: 300;
            }
            QFrame#dropZone[selected="true"] QLabel#dropIcon {
                font-size: 40px;
                font-weight: 800;
            }
            QLabel#dropTitle {
                color: #EDE9E1;
                font-size: 18px;
                font-weight: 700;
            }
            QFrame#dropZone[selected="true"] QLabel#dropTitle {
                color: #F4E4BF;
                font-size: 27px;
                font-weight: 800;
            }
            QLabel#dropSub {
                color: #68645D;
                font-size: 12px;
            }
            QFrame#fileBar {
                background: #121213;
                border: 1px solid #242321;
                border-radius: 10px;
            }
            QLabel#fileName {
                color: #C9C4BA;
                font-size: 12px;
            }
            QLabel#status {
                color: #817C73;
                font-size: 11px;
            }
            QPushButton#selectButton {
                background: #D6B16B;
                color: #11100E;
                border: none;
                border-radius: 9px;
                padding: 0 24px;
                font-weight: 800;
            }
            QPushButton#selectButton:hover {
                background: #E2C27C;
            }
            QFrame#conversionPanel {
                background: #131313;
                border: 1px solid #39342C;
                border-radius: 14px;
            }
            QFrame#sourceDetails {
                background: #181715;
                border: 1px solid #292722;
                border-radius: 10px;
            }
            QLabel#panelEyebrow {
                color: #9C8558;
                font-size: 10px;
                font-weight: 800;
                letter-spacing: 1px;
            }
            QLabel#sourceType {
                color: #D6B16B;
                font-size: 25px;
                font-weight: 800;
            }
            QLabel#sourceName {
                color: #F2E8D3;
                font-size: 15px;
                font-weight: 750;
            }
            QLabel#sourceMeta, QLabel#targetHint {
                color: #817C73;
                font-size: 11px;
            }
            QLabel#targetLabel {
                color: #F0ECE4;
                font-size: 13px;
                font-weight: 700;
            }
            QComboBox#targetFormat {
                background: #1A1917;
                color: #F2E8D3;
                border: 1px solid #655331;
                border-radius: 8px;
                padding: 8px 11px;
                min-height: 22px;
                font-weight: 650;
            }
            QComboBox#targetFormat::drop-down { border: none; width: 28px; }
            QComboBox#targetFormat QAbstractItemView {
                background: #1A1917;
                color: #F2E8D3;
                selection-background-color: #D6B16B;
                selection-color: #11100E;
            }
            QCheckBox#outputOption {
                color: #AAA49A;
                font-size: 11px;
                spacing: 8px;
            }
            QCheckBox#outputOption::indicator {
                width: 15px; height: 15px;
                border: 1px solid #655331;
                border-radius: 4px;
                background: #111110;
            }
            QCheckBox#outputOption::indicator:checked { background: #D6B16B; }
            QPushButton#convertButton {
                background: #D6B16B;
                color: #11100E;
                border: none;
                border-radius: 9px;
                min-height: 42px;
                font-size: 12px;
                font-weight: 850;
                padding: 0 24px;
            }
            QPushButton#convertButton:hover { background: #E5C980; }
            QPushButton#convertButton:disabled { background: #3C362B; color: #85795F; }
            QFrame#conversionCard {
                background: #141414;
                border: 1px solid #272522;
                border-radius: 13px;
            }
            QFrame#conversionCard:hover {
                border: 1px solid #80683E;
                background: #171614;
            }
            QFrame#conversionCard:disabled {
                background: #101010;
                border-color: #201F1D;
            }
            QFrame#conversionCard:disabled QLabel#cardIcon {
                color: #615846;
            }
            QFrame#conversionCard:disabled QLabel#cardTitle {
                color: #68645D;
            }
            QFrame#conversionCard:disabled QLabel#cardSubtitle {
                color: #4D4A45;
            }
            QLabel#cardIcon {
                color: #D6B16B;
                font-size: 25px;
                font-weight: 600;
            }
            QLabel#cardTitle {
                color: #EDE9E1;
                font-size: 14px;
                font-weight: 750;
            }
            QLabel#cardSubtitle {
                color: #716D66;
                font-size: 10px;
            }
            QPushButton#cardButton {
                background: #1C1B19;
                color: #B9A276;
                border: 1px solid #39342C;
                border-radius: 7px;
                font-size: 10px;
                font-weight: 750;
            }
            QPushButton#cardButton:hover {
                background: #D6B16B;
                color: #11100E;
            }
            QPushButton#cardButton:disabled {
                color: #4B4843;
                border-color: #242320;
                background: #151514;
            }
            QProgressBar {
                background: #181817;
                border: none;
                border-radius: 5px;
                height: 8px;
            }
            QProgressBar::chunk {
                background: #D6B16B;
                border-radius: 5px;
            }
            QMessageBox {
                background: #0F0F10;
            }
            QMessageBox QLabel {
                color: #D6B16B;
                font-size: 13px;
                min-width: 340px;
            }
            QMessageBox QPushButton {
                background: #1B1916;
                color: #D6B16B;
                border: 1px solid #6B5633;
                border-radius: 7px;
                min-width: 90px;
                min-height: 32px;
                font-weight: 750;
            }
            QMessageBox QPushButton:hover {
                background: #D6B16B;
                color: #11100E;
            }
            QScrollBar:vertical {
                background: #0D0D0E;
                width: 8px;
            }
            QScrollBar::handle:vertical {
                background: #33312D;
                border-radius: 4px;
            }
        """
        self.setStyleSheet(self._base_stylesheet)

        central = QWidget()
        self.setCentralWidget(central)
        root = QHBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Sidebar
        self.sidebar = QFrame()
        self.sidebar.setObjectName("sidebar")
        self.sidebar.setMinimumWidth(205)
        side = QVBoxLayout(self.sidebar)
        side.setSizeConstraint(QLayout.SetMinimumSize)
        side.setContentsMargins(18, 24, 18, 18)
        side.setSpacing(8)

        self.brand_logo = QLabel()
        self.brand_logo.setObjectName("brandLogo")
        self.brand_logo.setAlignment(Qt.AlignCenter)
        self.brand_logo.setFixedSize(179, 115)
        side.addWidget(self.brand_logo, 0, Qt.AlignHCenter)

        self.emir_video_widget = QVideoWidget() if QT_MULTIMEDIA_AVAILABLE else QLabel()
        self.emir_video_widget.setObjectName("emirVideo")
        self.emir_video_widget.setFixedSize(179, 82)
        if QT_MULTIMEDIA_AVAILABLE:
            # Siyah şerit bırakmadan tüm alanı kapla.
            self.emir_video_widget.setAspectRatioMode(Qt.IgnoreAspectRatio)
        self.emir_video_widget.setStyleSheet("border: 1px solid #5D252E; border-radius: 7px;")
        self.emir_video_widget.hide()
        side.addWidget(self.emir_video_widget, 0, Qt.AlignHCenter)

        self.emir_photo = QLabel()
        self.emir_photo.setObjectName("emirPhoto")
        self.emir_photo.setAlignment(Qt.AlignCenter)
        self.emir_photo.setFixedSize(82, 112)
        self.emir_photo.setStyleSheet("background: #0A090A; border: 1px solid #5D252E; border-radius: 8px;")
        self.emir_photo.hide()
        side.addWidget(self.emir_photo, 0, Qt.AlignHCenter)

        self.emir_audio = None
        self.emir_player = None
        if QT_MULTIMEDIA_AVAILABLE:
            self.emir_audio = QAudioOutput(self)
            self.emir_audio.setMuted(True)
            self.emir_player = QMediaPlayer(self)
            self.emir_player.setAudioOutput(self.emir_audio)
            self.emir_player.setVideoOutput(self.emir_video_widget)
            self.emir_player.setLoops(QMediaPlayer.Infinite)

        side.addSpacing(24)

        self.nav_converter = NavButton("Dönüştürücü", 0)
        self.nav_history = NavButton("Geçmiş", 1)
        self.nav_about = NavButton("Hakkında", 2)
        for b in [self.nav_converter, self.nav_history, self.nav_about]:
            b.setObjectName("nav")
            side.addWidget(b)

        # Sidebar navigation is active, not decorative.
        self.nav_converter.clicked.connect(self.show_converter)
        self.nav_history.clicked.connect(self.show_history)
        self.nav_about.clicked.connect(self.show_about)

        self.nav_converter.setChecked(True)
        side.addStretch(1)

        mode_panel = QFrame()
        mode_panel.setObjectName("modePanel")
        mode_layout = QVBoxLayout(mode_panel)
        mode_layout.setContentsMargins(5, 5, 5, 5)
        mode_layout.setSpacing(3)

        self.mode_toggle = QPushButton("MODLAR   ▾")
        self.mode_toggle.setObjectName("modeToggle")
        self.mode_toggle.setCheckable(True)
        self.mode_toggle.clicked.connect(self._toggle_mode_options)
        mode_layout.addWidget(self.mode_toggle)

        self.mode_options = QWidget()
        options_layout = QVBoxLayout(self.mode_options)
        options_layout.setContentsMargins(4, 3, 4, 4)
        options_layout.setSpacing(6)

        self.mode_group = QButtonGroup(self)
        self.mode_group.setExclusive(True)
        self.mode_buttons = {}
        for mode_key in ("azra", "rafine", "emir"):
            button = ModeButton(THEME_CONFIGS[mode_key]["name"], mode_key)
            button.clicked.connect(lambda _checked=False, key=mode_key: self.apply_theme(key))
            self.mode_group.addButton(button)
            self.mode_buttons[mode_key] = button
            options_layout.addWidget(button)
        self.mode_options.hide()
        mode_layout.addWidget(self.mode_options)
        side.addWidget(mode_panel)

        version = QLabel(f"{APP_NAME}\nv{APP_VERSION}")
        version.setObjectName("version")
        version.setAlignment(Qt.AlignCenter)
        side.addWidget(version)

        self.sidebar_scroll = QScrollArea()
        self.sidebar_scroll.setObjectName("sidebarScroll")
        self.sidebar_scroll.setWidgetResizable(True)
        self.sidebar_scroll.setFrameShape(QFrame.NoFrame)
        self.sidebar_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.sidebar_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.sidebar_scroll.setFixedWidth(225)
        self.sidebar_scroll.setWidget(self.sidebar)
        root.addWidget(self.sidebar_scroll)

        # Main content
        content = QWidget()
        content.setObjectName("mainContent")
        content.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        content_layout = QVBoxLayout(content)
        content_layout.setContentsMargins(42, 32, 42, 28)
        content_layout.setSpacing(15)

        self.emir_star = QLabel()
        self.emir_star.setObjectName("emirStar")
        self.emir_star.setAlignment(Qt.AlignCenter)
        self.emir_star.setFixedHeight(86)
        self.emir_star.hide()
        content_layout.addWidget(self.emir_star, 0, Qt.AlignHCenter)

        top = QHBoxLayout()
        heading = QVBoxLayout()
        heading.setSpacing(4)

        eyebrow = QLabel("DOCUMENT TOOLS")
        eyebrow.setObjectName("eyebrow")
        heading.addWidget(eyebrow)

        title = QLabel("Dosyalarınızı dönüştürün.")
        title.setObjectName("pageTitle")
        heading.addWidget(title)

        subtitle = QLabel("PDF, Word ve Excel arasında altı yönlü, yüksek kaliteli dönüşüm.")
        subtitle.setObjectName("pageSubtitle")
        heading.addWidget(subtitle)

        top.addLayout(heading)
        top.addStretch(1)

        content_layout.addLayout(top)

        self.drop_zone = DropZone()
        self.drop_zone.fileDropped.connect(self.set_file)
        self.drop_zone.clicked.connect(self.choose_file)
        content_layout.addWidget(self.drop_zone)

        self.conversion_panel = QFrame()
        self.conversion_panel.setObjectName("conversionPanel")
        panel_layout = QHBoxLayout(self.conversion_panel)
        panel_layout.setContentsMargins(18, 18, 18, 18)
        panel_layout.setSpacing(22)

        source_panel = QFrame()
        source_panel.setObjectName("sourceDetails")
        source_panel.setMinimumWidth(255)
        source_layout = QVBoxLayout(source_panel)
        source_layout.setContentsMargins(16, 15, 16, 15)
        source_layout.setSpacing(5)

        source_eyebrow = QLabel("SEÇİLEN DOSYA")
        source_eyebrow.setObjectName("panelEyebrow")
        source_layout.addWidget(source_eyebrow)

        self.source_type = QLabel()
        self.source_type.setObjectName("sourceType")
        source_layout.addWidget(self.source_type)

        self.source_name = QLabel()
        self.source_name.setObjectName("sourceName")
        self.source_name.setWordWrap(True)
        source_layout.addWidget(self.source_name)

        self.source_meta = QLabel()
        self.source_meta.setObjectName("sourceMeta")
        source_layout.addWidget(self.source_meta)
        source_layout.addStretch(1)

        change_file = QPushButton("DOSYAYI DEĞİŞTİR")
        change_file.setObjectName("cardButton")
        change_file.setMinimumHeight(32)
        change_file.clicked.connect(self.choose_file)
        source_layout.addWidget(change_file)
        panel_layout.addWidget(source_panel, 1)

        options = QVBoxLayout()
        options.setSpacing(8)
        target_eyebrow = QLabel("DÖNÜŞTÜRME AYARLARI")
        target_eyebrow.setObjectName("panelEyebrow")
        options.addWidget(target_eyebrow)

        target_label = QLabel("Hedef biçim")
        target_label.setObjectName("targetLabel")
        options.addWidget(target_label)

        self.target_format = QComboBox()
        self.target_format.setObjectName("targetFormat")
        self.target_format.currentIndexChanged.connect(self._target_format_changed)
        options.addWidget(self.target_format)

        self.target_hint = QLabel()
        self.target_hint.setObjectName("targetHint")
        self.target_hint.setWordWrap(True)
        options.addWidget(self.target_hint)

        self.pdf_excel_options = QWidget()
        pdf_excel_options_layout = QVBoxLayout(self.pdf_excel_options)
        pdf_excel_options_layout.setContentsMargins(0, 4, 0, 0)
        pdf_excel_options_layout.setSpacing(5)
        pdf_excel_label = QLabel("PDF sayfalarını Excel'de")
        pdf_excel_label.setObjectName("targetLabel")
        pdf_excel_options_layout.addWidget(pdf_excel_label)
        self.pdf_excel_page_mode = QComboBox()
        self.pdf_excel_page_mode.setObjectName("targetFormat")
        self.pdf_excel_page_mode.addItem(
            "Ayrı çalışma sayfaları", True
        )
        self.pdf_excel_page_mode.addItem(
            "Tek sayfada birleştir", False
        )
        pdf_excel_options_layout.addWidget(self.pdf_excel_page_mode)
        self.pdf_excel_options.hide()
        options.addWidget(self.pdf_excel_options)

        self.open_output_checkbox = QCheckBox("İşlem bitince çıktı klasörünü otomatik aç")
        self.open_output_checkbox.setObjectName("outputOption")
        options.addWidget(self.open_output_checkbox)
        options.addStretch(1)

        self.convert_button = QPushButton("DÖNÜŞTÜRMEYİ BAŞLAT")
        self.convert_button.setObjectName("convertButton")
        self.convert_button.clicked.connect(self.start_conversion)
        options.addWidget(self.convert_button)
        panel_layout.addLayout(options, 2)

        self.conversion_panel.setVisible(False)
        content_layout.addWidget(self.conversion_panel)

        self.progress = QProgressBar()
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        self.progress.setVisible(False)
        content_layout.addWidget(self.progress)

        self.status = QLabel("Hazır")
        self.status.setObjectName("status")
        self.status.setVisible(False)
        content_layout.addWidget(self.status)
        content_layout.addStretch(1)

        content_scroll = QScrollArea()
        content_scroll.setObjectName("contentScroll")
        content_scroll.setWidgetResizable(True)
        content_scroll.setFrameShape(QFrame.NoFrame)
        content_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        content_scroll.setWidget(content)
        root.addWidget(content_scroll, 1)

        self.update_buttons()
        saved_theme = self._window_settings.value("active_theme", "azra")
        if saved_theme not in THEME_CONFIGS:
            saved_theme = "azra"
        self.apply_theme(saved_theme)
        QTimer.singleShot(250, self._show_pending_update_result)

    @staticmethod
    def _set_scaled_pixmap(label, path):
        """Pikselleri Windows ekran ölçeklemesine göre üretir."""
        pixmap = QPixmap(path) if path else QPixmap()
        if pixmap.isNull():
            label.clear()
            return False
        ratio = label.devicePixelRatioF()
        pixel_size = QSize(
            round(label.width() * ratio),
            round(label.height() * ratio),
        )
        scaled = pixmap.scaled(
            pixel_size,
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation,
        )
        scaled.setDevicePixelRatio(ratio)
        label.setPixmap(scaled)
        return True

    def _stop_emir_media(self):
        if self.emir_player is not None:
            self.emir_player.stop()
            self.emir_player.setSource(QUrl())
        self.emir_video_widget.hide()
        self.emir_photo.hide()
        self.emir_star.hide()

    def _start_emir_media(self):
        photo_path = theme_asset_path("emir_photo")
        if self._set_scaled_pixmap(self.emir_photo, photo_path):
            self.emir_photo.show()

        star_path = theme_asset_path("emir_star")
        if self._set_scaled_pixmap(self.emir_star, star_path):
            self.emir_star.show()

        video_path = theme_asset_path("emir_video")
        if video_path and self.emir_player is not None:
            self.emir_player.setSource(QUrl.fromLocalFile(video_path))
            self.emir_video_widget.show()
            self.emir_player.play()

    def _toggle_mode_options(self, expanded):
        self.mode_options.setVisible(expanded)
        self.mode_toggle.setText("MODLAR   ▴" if expanded else "MODLAR   ▾")

    def apply_theme(self, mode_key):
        """Seçilen temayı uygular ve önceki moda ait özel medyayı temizler."""
        if mode_key not in THEME_CONFIGS:
            return

        self._stop_emir_media()
        config = THEME_CONFIGS[mode_key]
        self.active_theme = mode_key
        self.setStyleSheet(self._base_stylesheet + THEME_STYLESHEETS[mode_key])
        self.drop_zone.set_glow_colors(config["glow"], config["halo"])

        for button in (self.nav_converter, self.nav_history, self.nav_about):
            button.set_theme(mode_key)

        for key, button in self.mode_buttons.items():
            button.setChecked(key == mode_key)
        self.mode_toggle.setChecked(False)
        self._toggle_mode_options(False)

        logo_key = config.get("logo")
        if logo_key:
            self._set_scaled_pixmap(self.brand_logo, theme_asset_path(logo_key))
            self.brand_logo.show()
        else:
            self.brand_logo.clear()
            self.brand_logo.hide()

        if mode_key == "emir":
            self._start_emir_media()

        self._window_settings.setValue("active_theme", mode_key)

    def closeEvent(self, event):
        self._stop_emir_media()
        self._window_settings.setValue("window_geometry", self.saveGeometry())
        super().closeEvent(event)

    def _show_pending_update_result(self):
        result = consume_update_result()
        if result.get("status") == "success":
            version = clean_text(result.get("version")) or APP_VERSION
            QMessageBox.information(self, "Güncelleme tamamlandı", f"{APP_NAME} v{version} başarıyla yüklendi.")
        elif result.get("status") == "failed":
            detail = clean_text(result.get("message")) or "Bilinmeyen hata"
            QMessageBox.warning(
                self,
                "Güncelleme tamamlanamadı",
                "Güncelleme indirildi ancak dosyalar değiştirilemedi. "
                "Lütfen uygulamayı kapatıp güncellemeyi tekrar deneyin.\n\nAyrıntı: " + detail,
            )

    def _select_nav(self, active):
        for button in [self.nav_converter, self.nav_history,
                       self.nav_about]:
            button.blockSignals(True)
            button.setChecked(button is active)
            button.blockSignals(False)

    def show_converter(self):
        self._select_nav(self.nav_converter)
        self.status.setText("Hazır")

    def _history_path(self):
        return Path.home() / "AzraConverter_history.json"

    def _load_history(self):
        import json
        p = self._history_path()
        if not p.exists():
            return []
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return []

    def _save_history(self, mode, source, output, success=True):
        import json
        from datetime import datetime
        labels = {
            "pdf_excel": "PDF -> Excel",
            "pdf_word": "PDF -> Word",
            "excel_pdf": "Excel -> PDF",
            "word_pdf": "Word -> PDF",
            "word_excel": "Word -> Excel",
            "excel_word": "Excel -> Word",
        }
        data = self._load_history()
        data.insert(0, {
            "date": datetime.now().strftime("%d.%m.%Y %H:%M"),
            "file": Path(source).name,
            "operation": labels.get(mode, mode),
            "output": str(output),
            "status": "Başarılı" if success else "Hata",
        })
        self._history_path().write_text(
            json.dumps(data[:100], ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

    def _dark_dialog(self, title, width=760, height=500):
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.setMinimumSize(width, height)
        dialog.resize(width, height)
        dialog.setStyleSheet("""
            QDialog { background: #0F0F10; }
            QLabel { color: #E9E5DC; }
            QLabel#dialogTitle {
                color: #D6B16B; font-size: 22px; font-weight: 800;
            }
            QLabel#muted { color: #8E887E; font-size: 12px; }
            QTableWidget {
                background: #121212; color: #D9D4CB;
                gridline-color: #292824;
                border: 1px solid #2A2824; border-radius: 8px;
            }
            QHeaderView::section {
                background: #1B1A18; color: #D6B16B;
                border: none; padding: 9px; font-weight: 700;
            }
            QPushButton {
                background: #1A1917; color: #D6B16B;
                border: 1px solid #3B352B; border-radius: 8px;
                padding: 9px 20px; font-weight: 700;
            }
            QPushButton:hover {
                background: #D6B16B; color: #11100E;
            }
        """)
        return dialog

    def show_converter(self):
        self._select_nav(self.nav_converter)
        self.status.setText("Hazır")

    def show_history(self):
        self._select_nav(self.nav_history)

        # Modal dialog: exec() guarantees the history window is displayed.
        if hasattr(self, "_history_dialog") and self._history_dialog is not None:
            try:
                self._history_dialog.close()
            except Exception:
                pass

        dialog = self._dark_dialog(f"{APP_NAME} - Geçmiş", 820, 500)
        self._history_dialog = dialog
        dialog.finished.connect(lambda _=0: self._history_dialog_closed())

        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(28, 24, 28, 22)
        layout.setSpacing(12)

        title = QLabel("Dönüştürme Geçmişi")
        title.setObjectName("dialogTitle")
        layout.addWidget(title)

        subtitle = QLabel("Son işlemleriniz burada tutulur.")
        subtitle.setObjectName("muted")
        layout.addWidget(subtitle)

        table = QTableWidget()
        table.setColumnCount(5)
        table.setHorizontalHeaderLabels(["Tarih", "Dosya", "İşlem", "Durum", "Çıktı"])
        table.horizontalHeader().setStretchLastSection(True)
        table.setAlternatingRowColors(False)
        table.setEditTriggers(QTableWidget.NoEditTriggers)
        table.setSelectionBehavior(QTableWidget.SelectRows)

        history = self._load_history()
        table.setRowCount(len(history))

        for row, item in enumerate(history):
            values = [
                item.get("date", ""),
                item.get("file", ""),
                item.get("operation", ""),
                item.get("status", ""),
                Path(item.get("output", "")).name if item.get("output") else "",
            ]
            for col, value in enumerate(values):
                table.setItem(row, col, QTableWidgetItem(str(value)))

        layout.addWidget(table, 1)

        close = QPushButton("KAPAT")
        close.clicked.connect(dialog.close)
        layout.addWidget(close, 0, Qt.AlignRight)

        dialog.exec()
        self._history_dialog = None
        self._select_nav(self.nav_converter)

    def _history_dialog_closed(self):
        self._history_dialog = None
        self._select_nav(self.nav_converter)

    def check_for_updates(self):
        if getattr(self, "_update_thread", None) is not None:
            return

        self.update_status.setText("Güncellemeler kontrol ediliyor...")
        self.update_button.setEnabled(False)
        self._update_thread = QThread()
        self._update_worker = UpdateWorker()
        self._update_worker.moveToThread(self._update_thread)
        self._update_thread.started.connect(self._update_worker.run)
        self._update_worker.finished.connect(self.update_check_finished)
        self._update_worker.error.connect(self.update_check_error)
        self._update_worker.finished.connect(self._update_thread.quit)
        self._update_worker.error.connect(self._update_thread.quit)
        self._update_thread.finished.connect(self._update_worker.deleteLater)
        self._update_thread.finished.connect(self._update_thread_finished)
        self._update_thread.start()

    def _update_thread_finished(self):
        self._update_thread = None
        self._update_worker = None

    def _update_download_thread_finished(self):
        self._update_download_thread = None
        self._update_download_worker = None

    def update_check_finished(self, result):
        latest = result["version"]
        package_url = result.get("package_url", "")
        installer_url = result.get("download_url", "")
        if result["is_new"] and (installer_url or package_url):
            # Inno Setup updates the installed folder transactionally and handles
            # elevation itself. This is more reliable than replacing a live
            # PyInstaller directory with robocopy.
            self._update_kind = "installer" if installer_url else "package"
            self._update_download_url = installer_url or package_url
            self._update_version = latest
            self._update_checksum = (
                result.get("sha256", "")
                if installer_url else result.get("package_sha256", "")
            )
            note = f" - {result['notes']}" if result["notes"] else ""
            self.update_status.setText(f"Yeni sürüm hazır: v{latest}{note}")
            self.update_button.setText(
                "YENİ SÜRÜMÜ YÜKLE"
            )
            self.update_button.setEnabled(True)
            try:
                self.update_button.clicked.disconnect()
            except RuntimeError:
                pass
            self.update_button.clicked.connect(self.download_latest_update)
        elif result["is_new"]:
            self.update_status.setText(f"v{latest} hazır; indirme bağlantısı tanımlı değil.")
            self.update_button.setEnabled(True)
        else:
            self.update_status.setText(f"Güncel sürümü kullanıyorsunuz (v{APP_VERSION}).")
            self.update_button.setText("GÜNCELLEMELERİ KONTROL ET")
            self.update_button.setEnabled(True)

    def update_check_error(self, message):
        self.update_status.setText(message)
        self.update_button.setText("GÜNCELLEMELERİ KONTROL ET")
        self.update_button.setEnabled(True)

    def download_latest_update(self):
        url = getattr(self, "_update_download_url", "")
        version = getattr(self, "_update_version", "güncel")
        if not url or getattr(self, "_update_download_thread", None) is not None:
            return

        self.update_button.setEnabled(False)
        self.update_status.setText("Güncelleme indiriliyor... %0")
        self._update_download_thread = QThread()
        self._update_download_worker = UpdateDownloadWorker(
            url,
            version,
            getattr(self, "_update_checksum", ""),
        )
        self._update_download_worker.moveToThread(self._update_download_thread)
        self._update_download_thread.started.connect(self._update_download_worker.run)
        self._update_download_worker.progress.connect(self.update_download_progress)
        self._update_download_worker.finished.connect(self.update_download_finished)
        self._update_download_worker.error.connect(self.update_download_error)
        self._update_download_worker.finished.connect(self._update_download_thread.quit)
        self._update_download_worker.error.connect(self._update_download_thread.quit)
        self._update_download_thread.finished.connect(self._update_download_worker.deleteLater)
        self._update_download_thread.finished.connect(self._update_download_thread_finished)
        self._update_download_thread.start()

    def update_download_progress(self, progress):
        self.update_status.setText(f"Güncelleme indiriliyor... %{progress}")

    def update_download_finished(self, downloaded_path):
        self.update_status.setText("Güncelleme uygulanıyor. Program yeniden başlatılacak...")
        self.update_button.setText("GÜNCELLEME UYGULANIYOR")
        try:
            if getattr(self, "_update_kind", "installer") == "package":
                self._launch_in_app_update(downloaded_path)
            else:
                self._launch_update_installer(downloaded_path)
        except Exception as exc:
            self.update_download_error(str(exc))

    def update_download_error(self, message):
        self.update_status.setText(f"Güncelleme uygulanamadı: {message}")
        self.update_button.setText("YENİ SÜRÜMÜ YÜKLE")
        self.update_button.setEnabled(True)

    def _launch_update_installer(self, installer_path):
        installer = Path(installer_path)
        app_executable = Path(sys.executable).resolve()
        helper = Path(tempfile.gettempdir()) / "AzraConverterUpdates" / "install_update.cmd"
        helper.write_text(
            "@echo off\r\n"
            f"powershell -NoProfile -NonInteractive -Command \"Wait-Process -Id {os.getpid()} -ErrorAction SilentlyContinue\"\r\n"
            f"start \"\" /wait \"{installer}\" /SP- /VERYSILENT /SUPPRESSMSGBOXES /NORESTART\r\n"
            f"start \"\" \"{app_executable}\"\r\n"
            "del \"%~f0\"\r\n",
            encoding="utf-8",
        )
        creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
        subprocess.Popen(["cmd", "/c", str(helper)], creationflags=creationflags)
        QTimer.singleShot(350, QApplication.instance().quit)

    def _launch_in_app_update(self, package_path):
        """ZIP paketini doğrulanmış staging alanından uygulama klasörüne geçirir."""
        import ctypes

        if not getattr(sys, "frozen", False):
            raise RuntimeError("Program içi güncelleme yalnızca kurulu uygulamada çalışır.")

        update_root = Path(tempfile.mkdtemp(prefix="AzraConverterUpdate-"))
        staged_app = extract_update_package(package_path, update_root / "payload")
        current_executable = Path(sys.executable).resolve()
        install_dir = current_executable.parent
        helper = update_root / "apply_update.ps1"
        write_update_result("pending", self._update_version)
        result_path = update_result_path()
        helper.write_text(
            "$ErrorActionPreference = 'Stop'\r\n"
            f"$resultPath = '{_ps_quote(result_path)}'\r\n"
            f"$version = '{_ps_quote(self._update_version)}'\r\n"
            "try {\r\n"
            f"  Wait-Process -Id {os.getpid()} -ErrorAction SilentlyContinue\r\n"
            f"  $source = '{_ps_quote(staged_app)}'\r\n"
            f"  $target = '{_ps_quote(install_dir)}'\r\n"
            f"  $exe = '{_ps_quote(current_executable)}'\r\n"
            "  $arguments = @($source, $target, '/E', '/COPY:DAT', '/R:5', '/W:1', '/NFL', '/NDL', '/NJH', '/NJS', '/NP')\r\n"
            "  $copy = Start-Process -FilePath 'robocopy.exe' -ArgumentList $arguments -Wait -PassThru -WindowStyle Hidden\r\n"
            "  if ($copy.ExitCode -gt 7) { throw \"Dosyalar kopyalanamadı (robocopy: $($copy.ExitCode)).\" }\r\n"
            "  @{status='success'; version=$version; message=''} | ConvertTo-Json -Compress | Set-Content -LiteralPath $resultPath -Encoding UTF8\r\n"
            "  Start-Process -FilePath $exe\r\n"
            f"  Remove-Item -LiteralPath '{_ps_quote(Path(package_path))}' -Force -ErrorAction SilentlyContinue\r\n"
            "  Remove-Item -LiteralPath $source -Recurse -Force -ErrorAction SilentlyContinue\r\n"
            "} catch {\r\n"
            "  @{status='failed'; version=$version; message=$_.Exception.Message} | ConvertTo-Json -Compress | Set-Content -LiteralPath $resultPath -Encoding UTF8\r\n"
            "}\r\n"
            "Remove-Item -LiteralPath $PSCommandPath -Force -ErrorAction SilentlyContinue\r\n",
            encoding="utf-8-sig",
        )

        parameters = f'-NoProfile -ExecutionPolicy Bypass -File "{helper}"'
        result = ctypes.windll.shell32.ShellExecuteW(
            None, "runas", "powershell.exe", parameters, str(update_root), 0
        )
        if result <= 32:
            raise RuntimeError("Güncelleme için yönetici izni verilmedi.")
        QTimer.singleShot(350, QApplication.instance().quit)

    def show_about(self):
        self._select_nav(self.nav_about)

        dialog = self._dark_dialog(f"{APP_NAME} Hakkında", 480, 520)
        dialog.setMinimumSize(430, 450)
        dialog.setSizeGripEnabled(True)
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(26, 22, 26, 20)
        layout.setSpacing(7)

        # Azra Gold logosu yerine Türk bayrağı ve Emir Can Erarslan fotoğrafı.
        flag = QLabel()
        flag.setAlignment(Qt.AlignCenter)
        flag.setFixedSize(260, 138)
        self._set_scaled_pixmap(flag, theme_asset_path("turkish_flag"))
        layout.addWidget(flag, 0, Qt.AlignHCenter)

        photo = QLabel()
        photo.setAlignment(Qt.AlignCenter)
        photo.setFixedSize(72, 96)
        self._set_scaled_pixmap(photo, theme_asset_path("emir_photo"))
        layout.addWidget(photo, 0, Qt.AlignHCenter)

        brand = QLabel("CONVERTER")
        brand.setAlignment(Qt.AlignCenter)
        brand.setStyleSheet(
            "color:#D6B16B; font-size:14px; font-weight:800; letter-spacing:2px;"
        )
        layout.addWidget(brand)

        title = QLabel(APP_NAME)
        title.setObjectName("dialogTitle")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        version = QLabel(f"Sürüm {APP_VERSION}")
        version.setAlignment(Qt.AlignCenter)
        version.setStyleSheet("color:#8E887E; font-size:12px; font-weight:600;")
        layout.addWidget(version)

        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setStyleSheet("color:#2D2A25;")
        layout.addWidget(line)

        purpose_title = QLabel("PROGRAMIN AMACI")
        purpose_title.setAlignment(Qt.AlignCenter)
        purpose_title.setStyleSheet(
            "color:#D6B16B; font-size:12px; font-weight:800; letter-spacing:1px;"
        )
        layout.addWidget(purpose_title)

        purpose = QLabel(
            "PDF, Word ve Excel belgeleri arasında hızlı ve güvenilir dönüşüm sağlar.<br>"
            "Taranmış PDF'lerde OCR desteği sunar; belge düzenini mümkün olduğunca korur."
        )
        purpose.setAlignment(Qt.AlignCenter)
        purpose.setWordWrap(True)
        purpose.setStyleSheet("color:#C8C2B8; font-size:12px;")
        layout.addWidget(purpose)

        update_line = QFrame()
        update_line.setFrameShape(QFrame.HLine)
        update_line.setStyleSheet("color:#2D2A25;")
        layout.addWidget(update_line)

        update_title = QLabel("GÜNCELLEMELER")
        update_title.setAlignment(Qt.AlignCenter)
        update_title.setStyleSheet(
            "color:#D6B16B; font-size:12px; font-weight:800; letter-spacing:1px;"
        )
        layout.addWidget(update_title)

        self.update_status = QLabel("Güncellemeleri denetlemek için aşağıdaki düğmeyi kullanın.")
        self.update_status.setAlignment(Qt.AlignCenter)
        self.update_status.setWordWrap(True)
        self.update_status.setStyleSheet("color:#C8C2B8; font-size:11px;")
        layout.addWidget(self.update_status)

        self.update_button = QPushButton("GÜNCELLEMELERİ KONTROL ET")
        self.update_button.setMinimumHeight(36)
        self.update_button.clicked.connect(self.check_for_updates)
        layout.addWidget(self.update_button)

        layout.addStretch(1)

        signature = QLabel("Emir Can Erarslan")
        signature.setAlignment(Qt.AlignCenter)
        signature.setStyleSheet(
            "color:#D6B16B; font-size:16px; font-weight:700; font-style:italic;"
        )
        layout.addWidget(signature)

        signature_sub = QLabel(APP_NAME)
        signature_sub.setAlignment(Qt.AlignCenter)
        signature_sub.setStyleSheet("color:#777169; font-size:10px; letter-spacing:1px;")
        layout.addWidget(signature_sub)

        close = QPushButton("KAPAT")
        close.setMinimumHeight(38)
        close.clicked.connect(dialog.accept)
        layout.addWidget(close)

        dialog.exec()
        self.nav_converter.setChecked(True)

    def choose_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Dosya Seç",
            "",
            "Desteklenen Dosyalar (*.pdf *.doc *.docx *.docm *.dot *.dotx *.dotm *.odt *.rtf *.txt *.xls *.xlsx *.xlsm *.xlsb *.xlt *.xltx *.xltm *.ods *.csv *.tsv);;PDF (*.pdf);;Word (*.doc *.docx *.docm *.dot *.dotx *.dotm *.odt *.rtf *.txt);;Excel (*.xls *.xlsx *.xlsm *.xlsb *.xlt *.xltx *.xltm *.ods *.csv *.tsv)"
        )
        if path:
            self.set_file(path)

    def set_file(self, path):
        path = Path(path)
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            QMessageBox.warning(self, "Desteklenmeyen dosya",
                                "Bu dosya türü desteklenmiyor.\n\nPDF veya yaygın bir Word/Excel dosyası seçin.")
            return

        self.source_file = str(path)
        self.drop_zone.show_file(path)
        self.status.setText("Dosya hazır. Hedef biçimi seçip dönüştürmeyi başlatın.")
        self.status.setVisible(True)
        self.update_buttons()

    def _conversion_options_for_file(self):
        ext = Path(self.source_file).suffix.lower() if self.source_file else ""
        if ext in PDF_EXTENSIONS:
            return [
                ("pdf_excel", "Excel belgesi (.xlsx)", "OCR ile tabloları çalışma sayfasına aktarır."),
                ("pdf_word", "Word belgesi (.docx)", "OCR ile düzenlenebilir metin oluşturur."),
                ("pdf_delete_pages", "PDF sayfalarını sil (.pdf)", "Silinecek sayfaları seçip yeni bir PDF oluşturur."),
            ]
        if ext in SPREADSHEET_EXTENSIONS:
            return [
                ("excel_pdf", "PDF belgesi (.pdf)", "Çalışma sayfasını PDF olarak dışa aktarır."),
                ("excel_word", "Word belgesi (.docx)", "Metin ve tabloları Word belgesine aktarır."),
            ]
        if ext in WORD_EXTENSIONS:
            return [
                ("word_pdf", "PDF belgesi (.pdf)", "Belge düzenini koruyarak PDF oluşturur."),
                ("word_excel", "Excel belgesi (.xlsx)", "Metin ve tabloları çalışma sayfalarına aktarır."),
            ]
        return []

    def _target_format_changed(self):
        option = self.target_format.currentData()
        self.target_hint.setText(option[2] if option else "")
        self.pdf_excel_options.setVisible(bool(option and option[0] == "pdf_excel"))

    def update_buttons(self):
        options = self._conversion_options_for_file()
        self.conversion_panel.setVisible(bool(options))
        self.convert_button.setEnabled(bool(options) and self.thread is None)
        if not options:
            return

        selected_mode = self.target_format.currentData()
        selected_mode = selected_mode[0] if selected_mode else ""
        self.target_format.blockSignals(True)
        self.target_format.clear()
        for option in options:
            self.target_format.addItem(option[1], option)
        index = next((i for i, option in enumerate(options) if option[0] == selected_mode), 0)
        self.target_format.setCurrentIndex(index)
        self.target_format.blockSignals(False)
        self._target_format_changed()

        path = Path(self.source_file)
        try:
            size = path.stat().st_size
            size_text = f"{size / (1024 * 1024):.1f} MB" if size >= 1024 * 1024 else f"{max(1, size // 1024)} KB"
        except OSError:
            size_text = "Boyut bilgisi yok"
        self.source_type.setText(path.suffix.lstrip(".").upper() or "DOSYA")
        self.source_name.setText(path.name)
        self.source_meta.setText(f"{size_text}  •  Kaynak dosya")

    def start_conversion(self):
        if not self.source_file:
            QMessageBox.information(self, "Dosya seçin", "Önce bir dosya seçin.")
            return

        option = self.target_format.currentData()
        if not option:
            QMessageBox.warning(
                self, "Geçersiz dönüşüm",
                "Seçilen dosya bu dönüşüm türüyle uyumlu değil."
            )
            return
        mode = option[0]

        pages_to_remove = []
        if mode == "pdf_delete_pages":
            try:
                page_count = pdf_page_count(self.source_file)
            except Exception as exc:
                QMessageBox.critical(
                    self, "PDF okunamadı",
                    "PDF sayfaları okunamadı.\n\n" + (str(exc) or "Bilinmeyen hata"),
                )
                return
            if page_count < 2:
                QMessageBox.information(
                    self, "Silinecek sayfa yok",
                    "Sayfa silmek için PDF'te en az iki sayfa olmalı.",
                )
                return
            dialog = PdfPageDeletionDialog(page_count, self)
            if dialog.exec() != QDialog.DialogCode.Accepted:
                return
            pages_to_remove = dialog.selected_pages()

        self.progress.setVisible(False)
        self.progress.setValue(0)
        self.status.setText("Dönüştürülüyor...")
        self._active_mode = mode

        self.convert_button.setEnabled(False)

        self.thread = QThread()
        self.worker = ConverterWorker(
            mode,
            self.source_file,
            pdf_excel_separate_pages=bool(self.pdf_excel_page_mode.currentData()),
            pages_to_remove=pages_to_remove,
        )
        self.worker.moveToThread(self.thread)
        self.thread.started.connect(self.worker.run)
        self.worker.progress.connect(self.progress.setValue)
        self.worker.finished.connect(self.conversion_finished)
        self.worker.error.connect(self.conversion_error)
        self.worker.cancelled.connect(self.conversion_cancelled)
        self.worker.finished.connect(self.thread.quit)
        self.worker.error.connect(self.thread.quit)
        self.worker.cancelled.connect(self.thread.quit)
        self.thread.finished.connect(self.worker.deleteLater)
        self.thread.finished.connect(self.thread.deleteLater)
        self.thread.finished.connect(self.worker_thread_finished)

        self.progress_dialog = ConversionProgressDialog(
            Path(self.source_file).name, self
        )
        self.progress_dialog.cancelRequested.connect(self.cancel_conversion)
        self.worker.progress.connect(self.progress_dialog.set_progress)
        self.progress_dialog.show()
        self.progress_dialog.raise_()
        self.progress_dialog.activateWindow()
        self.thread.start()

    def cancel_conversion(self):
        if self.worker is not None:
            self.status.setText("Dönüştürme iptal ediliyor...")
            self.worker.cancel()

    def _close_progress_dialog(self):
        if self.progress_dialog is not None:
            self.progress_dialog.finish()
            self.progress_dialog.deleteLater()
            self.progress_dialog = None

    def worker_thread_finished(self):
        self.thread = None
        self.worker = None
        self.update_buttons()

    def _ask_open_output_folder(self, output):
        box = QMessageBox(self)
        box.setWindowTitle("İşlem tamamlandı")
        box.setIcon(QMessageBox.Information)
        box.setText("DÖNÜŞTÜRME TAMAMLANDI")
        box.setInformativeText(
            f"Dosya oluşturuldu: {Path(output).name}\n"
            "Çıktı klasörünü açmak ister misiniz?"
        )
        box.setStyleSheet("""
            QMessageBox { background: #0F0F10; }
            QMessageBox QLabel {
                color: #D6B16B; font-size: 13px;
                min-width: 320px; padding: 0;
            }
            QMessageBox QPushButton {
                background: #1B1916; color: #D6B16B;
                border: 1px solid #6B5633; border-radius: 8px;
                min-width: 120px; min-height: 36px; font-weight: 800;
            }
            QMessageBox QPushButton:hover {
                background: #D6B16B; color: #11100E;
            }
        """)
        open_button = box.addButton("KLASÖRÜ AÇ", QMessageBox.AcceptRole)
        box.addButton("KAPAT", QMessageBox.RejectRole)
        box.setDefaultButton(open_button)
        box.exec()
        return box.clickedButton() is open_button

    def conversion_finished(self, output):
        self._close_progress_dialog()
        self.progress.setValue(100)
        self.progress.setVisible(False)
        self.status.setText("Dönüştürme tamamlandı.")
        if self.source_file and self._active_mode:
            self._save_history(self._active_mode, self.source_file, output, True)
        self._active_mode = None
        if self.open_output_checkbox.isChecked() or self._ask_open_output_folder(output):
            os.startfile(str(Path(output).parent))

    def conversion_error(self, details):
        self._close_progress_dialog()
        if self.source_file and self._active_mode:
            self._save_history(self._active_mode, self.source_file, "", False)
        self._active_mode = None
        self.progress.setVisible(False)
        self.status.setText("Dönüştürme başarısız.")
        QMessageBox.critical(
            self, "Dönüştürme hatası",
            "İşlem sırasında hata oluştu.\n\n" +
            (details or "Bilinmeyen hata")
        )
        self.update_buttons()

    def conversion_cancelled(self):
        self._close_progress_dialog()
        self._active_mode = None
        self.progress.setVisible(False)
        self.progress.setValue(0)
        self.status.setText("Dönüştürme iptal edildi.")
        self.update_buttons()

def main():
    app = QApplication(sys.argv)
    app.setApplicationName(APP_NAME)
    app.setOrganizationName(APP_NAME)
    icon_path = theme_asset_path("app_icon")
    if os.path.exists(icon_path):
        app.setWindowIcon(QIcon(icon_path))

    font = QFont("Segoe UI", 10)
    app.setFont(font)

    settings = QSettings(APP_NAME, APP_NAME)
    saved_theme = settings.value("active_theme", "")
    onboarding_complete = settings.value("theme_onboarding_complete", False, type=bool)

    if saved_theme in THEME_CONFIGS:
        # Önceki sürümlerden gelen geçerli tema kaydı da tamamlanmış seçimdir.
        settings.setValue("theme_onboarding_complete", True)
    elif onboarding_complete:
        settings.setValue("active_theme", "azra")
    else:
        chooser = FirstRunThemeDialog()
        if chooser.exec() != QDialog.DialogCode.Accepted or chooser.selected_theme not in THEME_CONFIGS:
            return
        settings.setValue("active_theme", chooser.selected_theme)
        settings.setValue("theme_onboarding_complete", True)
    settings.sync()

    window = MainWindow()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
