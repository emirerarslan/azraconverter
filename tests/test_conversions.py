import sys
import tempfile
import types
import unittest
import zipfile
from pathlib import Path
from unittest import mock


class _DummySignal:
    def connect(self, *_args, **_kwargs):
        return None

    def emit(self, *_args, **_kwargs):
        return None


class _DummyWidget:
    def __init__(self, *_args, **_kwargs):
        pass


def _install_pyside_stubs():
    package = types.ModuleType("PySide6")
    core = types.ModuleType("PySide6.QtCore")
    gui = types.ModuleType("PySide6.QtGui")
    widgets = types.ModuleType("PySide6.QtWidgets")
    multimedia = types.ModuleType("PySide6.QtMultimedia")
    multimedia_widgets = types.ModuleType("PySide6.QtMultimediaWidgets")
    core.Qt = types.SimpleNamespace()
    core.QObject = _DummyWidget
    core.Signal = lambda *_args, **_kwargs: _DummySignal()
    core.QThread = core.QTimer = core.QSettings = core.QUrl = core.QSize = _DummyWidget
    core.QRectF = core.QPointF = _DummyWidget
    core.QPropertyAnimation = core.QEasingCurve = _DummyWidget
    for name in ("QFont", "QIcon", "QPixmap", "QPainter", "QPainterPath", "QColor", "QPen"):
        setattr(gui, name, _DummyWidget)
    for name in (
        "QApplication", "QMainWindow", "QWidget", "QVBoxLayout", "QHBoxLayout",
        "QGridLayout", "QLabel", "QPushButton", "QFileDialog", "QMessageBox",
        "QFrame", "QProgressBar", "QSizePolicy", "QSpacerItem", "QDialog",
        "QScrollArea", "QTableWidget", "QTableWidgetItem",
        "QComboBox", "QCheckBox", "QButtonGroup", "QLayout",
        "QGraphicsDropShadowEffect",
    ):
        setattr(widgets, name, _DummyWidget)
    multimedia.QAudioOutput = multimedia.QMediaPlayer = _DummyWidget
    multimedia_widgets.QVideoWidget = _DummyWidget
    sys.modules.update({
        "PySide6": package,
        "PySide6.QtCore": core,
        "PySide6.QtGui": gui,
        "PySide6.QtWidgets": widgets,
        "PySide6.QtMultimedia": multimedia,
        "PySide6.QtMultimediaWidgets": multimedia_widgets,
    })


_install_pyside_stubs()
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
import main  # noqa: E402


class Progress:
    def __init__(self):
        self.values = []

    def emit(self, value):
        self.values.append(value)


class ConversionTests(unittest.TestCase):
    def test_conversion_worker_honours_cancel_request(self):
        worker = main.ConverterWorker("pdf_word", "document.pdf")
        worker.cancel()
        with self.assertRaises(main.ConversionCancelled):
            worker.check_cancelled()

    def test_spreadsheet_compatibility_restores_missing_numpy_short_alias(self):
        try:
            import numpy
        except ImportError:
            self.skipTest("NumPy bu ortamda kurulu değil.")

        original = getattr(numpy, "short", None)
        if original is None:
            self.skipTest("Bu NumPy sürümünde kısa sayı türü tanımlı değil.")
        delattr(numpy, "short")
        try:
            main.ensure_spreadsheet_dependency_compatibility()
            self.assertIs(numpy.short, numpy.int16)
        finally:
            numpy.short = original

    def test_custom_update_manifest_is_normalised(self):
        manifest = main.normalise_update_manifest(
            {
                "version": "1.2.0",
                "download_url": "AZRA-SETUP.exe",
                "sha256": "abc123",
                "package_url": "AZRA-UPDATE.zip",
                "package_sha256": "def456",
                "notes": "Test sürümü",
            },
            "https://example.com/updates/version.json",
        )
        self.assertEqual(manifest["version"], "1.2.0")
        self.assertEqual(
            manifest["download_url"],
            "https://example.com/updates/AZRA-SETUP.exe",
        )
        self.assertEqual(manifest["sha256"], "abc123")
        self.assertEqual(
            manifest["package_url"],
            "https://example.com/updates/AZRA-UPDATE.zip",
        )
        self.assertEqual(manifest["package_sha256"], "def456")

    def test_github_release_api_manifest_is_normalised(self):
        manifest = main.normalise_update_manifest(
            {
                "tag_name": "v1.1.0",
                "body": "Yeni sürüm",
                "assets": [
                    {
                        "name": "AZRA-CONVERTER-SETUP-1.1.0.exe",
                        "browser_download_url": "https://example.com/setup.exe",
                        "digest": "sha256:abcdef",
                    },
                    {
                        "name": "AZRA-CONVERTER-UPDATE-1.1.1.zip",
                        "browser_download_url": "https://example.com/update.zip",
                        "digest": "sha256:123456",
                    },
                ],
            },
            "https://api.github.com/repos/example/app/releases/latest",
        )
        self.assertEqual(manifest["version"], "1.1.0")
        self.assertEqual(manifest["download_url"], "https://example.com/setup.exe")
        self.assertEqual(manifest["sha256"], "abcdef")
        self.assertEqual(manifest["package_url"], "https://example.com/update.zip")
        self.assertEqual(manifest["package_sha256"], "123456")

    def test_update_package_is_safely_extracted(self):
        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            package = folder / "update.zip"
            with zipfile.ZipFile(package, "w") as archive:
                archive.writestr("ConverteR.exe", b"binary")
                archive.writestr("_internal/library.dll", b"library")
            app_folder = main.extract_update_package(package, folder / "staging")
            self.assertEqual(app_folder, (folder / "staging").resolve())
            self.assertTrue((app_folder / "_internal" / "library.dll").exists())

    def test_update_package_rejects_path_traversal(self):
        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            package = folder / "unsafe.zip"
            with zipfile.ZipFile(package, "w") as archive:
                archive.writestr("../outside.txt", b"unsafe")
                archive.writestr("ConverteR.exe", b"binary")
            with self.assertRaises(RuntimeError):
                main.extract_update_package(package, folder / "staging")

    def test_update_urls_keep_fallbacks_without_duplicates(self):
        urls = main.update_manifest_urls({
            "manifest_urls": [main.DEFAULT_MANIFEST_URLS[0]],
            "manifest_url": main.DEFAULT_MANIFEST_URLS[2],
        })
        self.assertEqual(len(urls), len(set(urls)))
        self.assertIn("api.github.com", " ".join(urls))

    def test_update_result_is_consumed_once(self):
        with tempfile.TemporaryDirectory() as folder, mock.patch.object(
            main, "update_result_path", return_value=Path(folder) / "result.json"
        ):
            main.write_update_result("success", "1.1.6")
            self.assertEqual(main.consume_update_result(), {
                "status": "success", "version": "1.1.6", "message": "",
            })
            self.assertEqual(main.consume_update_result(), {})

    def test_extension_catalog_covers_legacy_and_open_formats(self):
        expected = {
            ".pdf", ".doc", ".docx", ".docm", ".odt", ".rtf", ".txt",
            ".xls", ".xlsx", ".xlsm", ".xlsb", ".ods", ".csv", ".tsv",
            ".mp4", ".avi", ".mov", ".mkv", ".webm", ".mpeg",
        }
        self.assertTrue(expected.issubset(main.SUPPORTED_EXTENSIONS))

    def test_video_conversion_requires_ffmpeg(self):
        with tempfile.TemporaryDirectory() as folder, mock.patch.object(
            main, "find_ffmpeg", return_value=None
        ):
            source = Path(folder) / "sample.mov"
            source.write_bytes(b"not-a-real-video")
            with self.assertRaisesRegex(RuntimeError, "FFmpeg bulunamadı"):
                main.video_convert(source, "mp4", Progress())

    def test_ffmpeg_progress_time_is_converted_to_seconds(self):
        self.assertAlmostEqual(main._ffmpeg_time_seconds("01:02:03.500000"), 3723.5)
        self.assertEqual(main._ffmpeg_time_seconds("geçersiz"), 0.0)

    def test_word_excel_and_excel_word_round_trip(self):
        from docx import Document
        from openpyxl import load_workbook

        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            source = folder / "örnek.docx"
            document = Document()
            document.add_heading("Satış Raporu", level=1)
            document.add_paragraph("Türkçe karakter testi: çğıöşü İ")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Ürün"
            table.cell(0, 1).text = "Tutar"
            table.cell(1, 0).text = "Altın"
            table.cell(1, 1).text = "1250"
            document.save(source)

            progress = Progress()
            excel_output = main.word_to_excel(source, progress)
            self.assertTrue(excel_output.exists())
            workbook = load_workbook(excel_output)
            self.assertIn("Belge Metni", workbook.sheetnames)
            self.assertIn("Tablo 1", workbook.sheetnames)
            self.assertEqual(workbook["Tablo 1"]["A2"].value, "Altın")

            word_output = main.excel_to_word(excel_output, progress)
            self.assertTrue(word_output.exists())
            converted = Document(word_output)
            self.assertGreaterEqual(len(converted.tables), 2)
            self.assertEqual(progress.values[-1], 100)

    def test_pdf_extraction_to_excel_and_word(self):
        from docx import Document
        from openpyxl import load_workbook
        from reportlab.pdfgen import canvas

        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            source = folder / "sample.pdf"
            pdf = canvas.Canvas(str(source))
            pdf.drawString(72, 760, "Product | Amount")
            pdf.drawString(72, 740, "Gold | 1250")
            pdf.save()

            progress = Progress()
            excel_output = main.pdf_to_excel(source, progress)
            workbook = load_workbook(excel_output)
            self.assertEqual(workbook.active["A1"].value, "Product")
            self.assertEqual(workbook.active["B2"].value, "1250")

            with mock.patch.object(
                main, "word_to_docx_with_microsoft_word",
                side_effect=RuntimeError("fallback test"),
            ):
                word_output = main.pdf_to_word(source, progress)
            converted = Document(word_output)
            self.assertIn("Product", "\n".join(p.text for p in converted.paragraphs))

    def test_pdf_excel_page_mode_can_merge_or_split_pages(self):
        from openpyxl import load_workbook
        from reportlab.pdfgen import canvas

        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            source = folder / "two-pages.pdf"
            pdf = canvas.Canvas(str(source))
            pdf.drawString(72, 760, "Gold | 1250")
            pdf.showPage()
            pdf.drawString(72, 760, "Silver | 750")
            pdf.save()

            split_output = main.pdf_to_excel(source, Progress(), separate_pages=True)
            split_workbook = load_workbook(split_output)
            self.assertEqual(split_workbook.sheetnames, ["Sayfa 1", "Sayfa 2"])

            merged_output = main.pdf_to_excel(source, Progress(), separate_pages=False)
            merged_workbook = load_workbook(merged_output)
            self.assertEqual(merged_workbook.sheetnames, ["Veriler"])
            values = [cell.value for row in merged_workbook.active.iter_rows() for cell in row]
            self.assertIn("Gold", values)
            self.assertIn("Silver", values)

    def test_pdf_page_removal_preserves_unselected_pages(self):
        from pypdf import PdfReader
        from reportlab.pdfgen import canvas

        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            source = folder / "three-pages.pdf"
            pdf = canvas.Canvas(str(source))
            for number in range(1, 4):
                pdf.drawString(72, 760, f"Page {number}")
                pdf.showPage()
            pdf.save()

            self.assertEqual(main.pdf_page_count(source), 3)
            output = main.remove_pdf_pages(source, [2], Progress())
            reader = PdfReader(output)
            self.assertEqual(len(reader.pages), 2)
            content = "\n".join(page.extract_text() or "" for page in reader.pages)
            self.assertIn("Page 1", content)
            self.assertIn("Page 3", content)
            self.assertNotIn("Page 2", content)

    def test_pdf_exports_are_valid(self):
        from docx import Document
        from openpyxl import Workbook
        from pypdf import PdfReader

        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            progress = Progress()

            word_source = folder / "source.docx"
            document = Document()
            document.add_heading("Başlık", level=1)
            document.add_paragraph("PDF kalite doğrulaması")
            document.save(word_source)
            word_pdf = main.word_to_pdf(word_source, progress)
            self.assertGreaterEqual(len(PdfReader(word_pdf).pages), 1)

            excel_source = folder / "source.xlsx"
            workbook = Workbook()
            workbook.active.append(["Ürün", "Tutar"])
            workbook.active.append(["Altın", 1250])
            workbook.save(excel_source)
            excel_pdf = main.excel_to_pdf(excel_source, progress)
            self.assertGreaterEqual(len(PdfReader(excel_pdf).pages), 1)

    def test_turkish_semicolon_csv_to_word(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as folder:
            folder = Path(folder)
            source = folder / "veriler.csv"
            source.write_bytes("Ürün;Tutar\r\nÇeyrek;1250\r\n".encode("cp1254"))
            output = main.excel_to_word(source, Progress())
            document = Document(output)
            self.assertEqual(document.tables[0].cell(1, 0).text, "Çeyrek")
            self.assertEqual(document.tables[0].cell(1, 1).text, "1250")


if __name__ == "__main__":
    unittest.main()
